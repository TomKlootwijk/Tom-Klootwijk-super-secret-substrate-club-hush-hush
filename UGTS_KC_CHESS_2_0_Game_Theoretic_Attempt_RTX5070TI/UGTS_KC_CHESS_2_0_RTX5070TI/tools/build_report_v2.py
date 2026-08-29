from __future__ import annotations

import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

# Reuse the visually-tested low-level helpers and style pack from the 1.0 report.
from build_report import (
    NAVY, INK, BLUE, CYAN, TEAL, GOLD, ORANGE, MAGENTA, RED, PURPLE,
    LIGHT, MID, GRAY, DARKGRAY, WHITE,
    make_styles, set_paragraph_bottom_border, add_page_field,
    set_cell_shading, set_cell_border, set_cell_margins,
    add_pic, caption, add_callout, add_table, add_bullets, add_numbered,
    add_code, add_section_title, add_source_note,
)

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
OUT = REPORT / "UGTS_KC_CHESS_2_0_Game_Theoretic_Campaign_RTX5070Ti.docx"

validation = json.loads((ROOT / "validation" / "summary_v2.json").read_text(encoding="utf-8"))
source_reg = json.loads((ROOT / "spec" / "source_register.json").read_text(encoding="utf-8"))
mechanisms = json.loads((ROOT / "spec" / "chess_mechanisms.json").read_text(encoding="utf-8"))
kq = json.loads((ROOT / "data" / "kqk.tb.json").read_text(encoding="utf-8"))
kr = json.loads((ROOT / "data" / "krk.tb.json").read_text(encoding="utf-8"))
workload = json.loads((ROOT / "examples" / "campaign" / "initial_depth4_workloads.json").read_text(encoding="utf-8"))
profile = json.loads((ROOT / "spec" / "rtx5070ti_profile.json").read_text(encoding="utf-8"))
mate_wdl = json.loads((ROOT / "examples" / "campaign" / "bounded_wdl_mate_over_claim.json").read_text(encoding="utf-8"))
initial_wdl = json.loads((ROOT / "examples" / "campaign" / "bounded_wdl_initial_depth2.json").read_text(encoding="utf-8"))

CANONICAL_ID = "ugts.application.chess-proof@2.0.0"
RULE_PROFILE = "fide-classical-2023-claims-as-actions-v2"
ROOT_FEN = workload["root_fen"]
ROOT_STATE_HASH = "eef9110182efec68661804d99fb8a91a1eac10973bdb4a44569d66ac0c9f8078"


def setup_document_v2() -> Document:
    doc = Document()
    make_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(15)
    sec.bottom_margin = Mm(15)
    sec.left_margin = Mm(17)
    sec.right_margin = Mm(17)
    sec.header_distance = Mm(6)
    sec.footer_distance = Mm(7)
    sec.different_first_page_header_footer = True

    header = sec.header
    p = header.paragraphs[0]
    p.text = "UGTS-KC CHESS 2.0  |  CLASSICAL GAME-THEORETIC PROOF CAMPAIGN"
    p.style = doc.styles["Small"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_bottom_border(p, color=MID, size="6")

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Prepared for Tom Klootwijk  |  Version 2.0.0  |  Page ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(p)

    props = doc.core_properties
    props.title = "UGTS-KC Chess 2.0 - Classical Game-Theoretic Proof Campaign"
    props.subject = "Proof-preserving classical chess campaign and RTX 5070 Ti Laptop CUDA handoff"
    props.author = "OpenAI, prepared for Tom Klootwijk"
    props.keywords = "UGTS, chess, game theory, proof certificate, WDL, CUDA, RTX 5070 Ti Laptop"
    props.comments = "Requester attribution recorded as supplied and not independently verified."
    return doc


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (a, b) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        for c in (c1, c2):
            set_cell_margins(c, top=45, start=80, bottom=45, end=80)
            border = {"val": "single", "sz": "4", "color": MID}
            set_cell_border(c, top=border, bottom=border, left=border, right=border)
        set_cell_shading(c1, "EAF3FA")
        set_cell_shading(c2, WHITE)
        c1.width = Inches(1.65)
        c2.width = Inches(4.95)
        r = c1.paragraphs[0].add_run(a)
        r.bold = True
        r.font.size = Pt(8.4)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r = c2.paragraphs[0].add_run(b)
        r.font.size = Pt(8.4)
        r.font.color.rgb = RGBColor.from_string(INK)


def cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Cover Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("UGTS-KC APPLICATION / PROOF PROFILE")

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("UGTS-KC CHESS 2.0")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Classical Game-Theoretic Proof Campaign")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Exact action semantics, four-valued WDL certificates, root-sharded campaign ledger\nand an RTX 5070 Ti Laptop CUDA execution handoff")
    r.font.size = Pt(11.2)
    r.font.color.rgb = RGBColor.from_string(DARKGRAY)

    add_pic(doc, "cover_architecture.png", width=6.8,
            alt="Architecture of the UGTS Chess 2.0 proof campaign showing exact state and actions, WDL obligations, proof shards, lineage, Python authority, C++ host code and optional CUDA proposal kernels.")

    add_metadata_table(doc, [
        ("Prepared for", "Tom Klootwijk"),
        ("Canonical identity", CANONICAL_ID),
        ("Document date", "29 August 2026"),
        ("Target profile", "GeForce RTX 5070 Ti Laptop GPU, 12 GB GDDR7, SM 12.0 build profile"),
        ("Captured root value", "UNKNOWN - the orthodox initial position is not claimed solved"),
    ])
    add_callout(doc, "Evidence boundary",
                "This deliverable is an attempted game-theoretic solution framework, not a declaration that classical chess has been solved. The package promotes only independently checked exact rules, certificates and closed finite graphs. Any open edge, history ambiguity, draw uncertainty or budget stop remains UNKNOWN.",
                fill="FFF4F2", accent=RED)
    doc.add_page_break()


def executive(doc: Document) -> None:
    add_section_title(doc, "0", "Executive Decision and Captured Result",
                      "The release upgrades the earlier bounded solver into a proof campaign that Codex can build, measure and extend on the named laptop.")
    add_callout(doc, "Release decision - GO",
                "Deliver the exact state/action foundation, independent proof semantics, twenty root obligations, portable campaign ledger, C++20 host executables, optional CUDA kernels and SM120 device profile. Do not coerce the unresolved initial position into WIN, DRAW or LOSS.",
                fill="EAF7F3", accent=TEAL)

    rows = [
        ("Orthodox root", "UNKNOWN", "No child root obligation has an independently verified WDL certificate."),
        ("Exact rules", "PASS", "Strict FEN, legal moves, castling, en-passant, promotion, king safety and terminal semantics."),
        ("Bounded proofs", "PASS", "Mate certificate and WDL certificate examples independently replay and verify."),
        ("Finite exact domains", "PASS", "Bundled KQK/KRK WDL plus DTM tables; maxima 20 and 32 plies."),
        ("Host implementation", "PASS", "Python package, C++20 binaries, CPU packed fallback and deterministic campaign files."),
        ("Physical RTX run", "PENDING", "No nvcc or RTX 5070 Ti Laptop GPU was available in the packaging environment."),
    ]
    add_table(doc, ["Layer", "Status", "Meaning"], rows, widths=[1.45, 0.75, 4.35], font_size=8.0, first_col_bold=True)

    doc.add_heading("What is delivered", level=2)
    add_bullets(doc, [
        "A complete source ZIP that Codex can build on Windows, Linux or WSL, with a dedicated RTX 5070 Ti Laptop CMake preset.",
        "A PDF/DOCX foundation separating exact game semantics, exact certificates, heuristic search and unfinished global proof work.",
        "A root campaign database and twenty content-addressed first-move obligations; exact depth-four workloads sum to 197,281 legal leaf paths.",
        "A checker-first GPU protocol: CUDA may accelerate candidate expansion and monotone fixed-point proposals, but the host verifier remains authoritative.",
    ], compact=False)

    add_callout(doc, "Captured validation",
                f"{validation['python_tests']['count']} Python tests; {validation['python_perft']['checks']} Python perft checks; {validation['native_build']['ctest_total']}/3 native CTest cases; {validation['packed_differential']['verified_moves']} packed moves with zero mismatches; {validation['schemas_and_proofs']['schema_documents']} schema instances; clean-install wheel built.",
                fill="EEF4FF", accent=BLUE)
    doc.add_page_break()


def source_grounding(doc: Document) -> None:
    add_section_title(doc, "1", "Source Grounding and Component Identity",
                      "The UGTS documents provide the authority and evidence discipline; the chess rules, WDL calculus and CUDA implementation are visibly engineering-derived.")
    add_pic(doc, "source_grounding.png", width=6.75,
            alt="Source grounding diagram connecting UGTS GPU-native, foundation, runtime, literal, SCLP, Elizabeth, Operator 4.2 and Chess 1.0 sources to the Chess 2.0 engineering delta.")
    caption(doc, "Source lineage used by this release. The wallet-specific SARA profile was reviewed and excluded from chess mechanics.")

    rows = [
        ("UGTS-GN 1.1", "Canonical support -> compatibility -> guard -> event -> transition -> lineage; packed precision boundaries."),
        ("UGTS-KC 2.0", "Typed state, topology, bounded queries, dynamics and explicit event-degeneracy handling."),
        ("Two Hands 3.0", "Proposal verification, deterministic commit, pre/post hashes, checkpoints and replay."),
        ("Literal 3.6 / SCLP 3.6.2", "Content-addressed definitions, dependency closure, finite keys, bounded branching and storage audits."),
        ("Elizabeth 3.9 / Operator 4.2", "Deterministic runtime and explicit parse, dependency, evaluation, event and replay order."),
        ("Chess 1.0", "Immediate exact rule, mate-certificate and KQK/KRK baseline."),
    ]
    add_table(doc, ["Source", "Retained contribution"], rows, widths=[1.65, 4.95], font_size=7.7, first_col_bold=True)

    p = doc.add_paragraph()
    p.add_run("Governed identity. ").bold = True
    p.add_run(f"The canonical release identity is {CANONICAL_ID}. This follows the versioning charter's rule that component-scoped identities and explicit parent relationships are more truthful than a single global number ladder.")
    add_source_note(doc, "UGTS-GN 1.1, UGTS-KC 2.0, Two Hands 3.0, Literal 3.6, SCLP 3.6.2, Elizabeth 3.9, General Operator and Order 4.2, Chess 1.0 and the Phase 1 Versioning Charter are registered in spec/source_register.json with exact SHA-256 hashes.")
    doc.add_page_break()


def solve_meanings(doc: Document) -> None:
    add_section_title(doc, "2", "What 'Solve Classical Chess' Means",
                      "A precise campaign must distinguish rule correctness, solved finite positions and a solved initial game.")
    add_pic(doc, "solve_ladder.png", width=6.75,
            alt="Three-stage ladder distinguishing rules solved, positions solved, and the initial game solved, with the final stage marked not established and root unknown.")
    caption(doc, "Three proof scopes. The package closes the first and selected cases of the second; it does not close the third.")

    rows = [
        ("Rule exactness", "Legal actions and deterministic transitions are reproducible from the same state and rule profile."),
        ("Certificate exactness", "Every quantified WIN/LOSS/DRAW obligation is represented and independently checked."),
        ("Tablebase exactness", "Every legal state in a declared finite material partition has an exact WDL value and, where decisive, DTM."),
        ("Search estimate", "A score, principal variation, proof-number priority or finite cutoff guides work but is not itself a theorem."),
        ("Unknown", "A safe and meaningful result whenever required edges, histories, draw logic or verifier evidence remain open."),
    ]
    add_table(doc, ["Class", "Meaning"], rows, widths=[1.55, 5.05], font_size=8.2, first_col_bold=True)
    add_callout(doc, "Non-claim",
                "This report does not claim a weak, strong or ultra-weak solution of the 32-piece initial position. It specifies a falsifiable route to such a result and records the current root as UNKNOWN.",
                fill="FFF4F2", accent=RED)
    doc.add_page_break()


def state_profile(doc: Document) -> None:
    add_section_title(doc, "3", "Classical State, History and Identity",
                      "A board diagram or FEN alone is not a complete proof state when repetition and move-count rules affect the result.")
    add_pic(doc, "state_identity.png", width=6.75,
            alt="State identity diagram separating board, side, castling, legal en-passant right, halfmove and fullmove fields and showing distinct position and game-state hashes.")
    caption(doc, "The serialized position hash and semantic game-state hash serve different purposes.")

    add_code(doc, "Q = (board, side_to_move, castling_rights, legal_en_passant_right,\n     halfmove_clock, repetition_count_context, rule_profile, lineage)\n\nserialized FEN identity != repetition identity != proof-state identity")
    add_bullets(doc, [
        "Repetition identity includes an en-passant field only when a legal en-passant capture is available, matching the move-possibility basis of the rule.",
        "The semantic proof-state hash excludes the non-semantic fullmove display number but includes the halfmove clock and exact sorted repetition counts.",
        "A 64-bit transposition key may index a cache. It is never accepted as the complete proof identity or as collision-free lineage.",
        "Coordinates identify squares, not history. Two visually identical boards can have different castling, en-passant, claim and repetition semantics.",
    ], compact=False)
    add_callout(doc, "Root identity", f"Initial FEN: {ROOT_FEN}\nSemantic root SHA-256: {ROOT_STATE_HASH}", fill="EEF4FF", accent=BLUE)
    doc.add_page_break()


def draw_semantics(doc: Document) -> None:
    add_section_title(doc, "4", "Draw Claims, Automatic Terminals and Dead Position",
                      "The rules profile treats player claims as actions and automatic draw conditions as terminals, with checkmate priority preserved.")
    rows = [
        ("Threefold repetition", "Optional action", "Current position or a declared intended legal move can establish the required occurrence."),
        ("50-move rule", "Optional action", "Current clock threshold or an intended legal move can create the threshold."),
        ("Fivefold repetition", "Automatic draw", "Terminal without a player claim."),
        ("75-move rule", "Automatic draw", "Terminal; checkmate on the last move takes precedence."),
        ("Stalemate", "Automatic draw", "No legal move and king not in check."),
        ("Dead position", "Automatic draw", "Neither player can checkmate by any possible series of legal moves."),
        ("Draw agreement / clock / arbiter", "Outside game-value core", "Tournament-contingent procedures are not treated as intrinsic board-state transitions."),
    ]
    add_table(doc, ["Condition", "Model", "Exact treatment"], rows, widths=[1.5, 1.2, 3.9], font_size=7.5, first_col_bold=True)

    doc.add_heading("Claim-as-action consequence", level=2)
    p = doc.add_paragraph()
    p.add_run("A claimable draw is not forced merely because it is available. ").bold = True
    p.add_run("A winning move may be chosen instead. The bundled example starts with a valid current 50-move claim and still proves WIN because Qa4# immediately checkmates.")
    add_code(doc, "available_actions(q) = legal_moves(q) U valid_draw_claims(q)\nterminal(q) checks: checkmate -> stalemate/dead -> fivefold/75-move\ncheckmate outranks the automatic 75-move terminal on the final move")
    add_callout(doc, "Current limitation",
                "The built-in dead-position recognizer certifies a conservative exact subset. A global solution requires a complete dead-position certificate/oracle or must keep affected nodes UNKNOWN.",
                fill="FFF9E8", accent=GOLD)
    add_source_note(doc, "FIDE Laws of Chess effective 1 January 2023, Articles 5.2.1-5.2.2 and 9.2-9.6. The exact implementation profile is named in every certificate.")
    doc.add_page_break()


def move_authority(doc: Document) -> None:
    add_section_title(doc, "5", "Canonical Move Authority and Operation Order",
                      "Candidate generation may be parallel or heuristic. State mutation and proof promotion are deterministic and ordered.")
    add_pic(doc, "authority_chain.png", width=6.75,
            alt="Six-stage UGTS authority chain for a chess move: support, compatibility, special-rule readiness, king-safety verification, atomic transition and history/hash lineage.")
    caption(doc, "Chess-specific mechanics mapped into the permanent UGTS authority sequence.")

    phases = [
        ("0-2", "parse / normalize / resolve", "FEN, moves, rules, definitions and profile references"),
        ("3-5", "type / canonicalize / plan", "Board, history, state hashes, ordered candidate set"),
        ("6-7", "evaluate / certify", "Pseudo moves, attacks, castling/EP/promotion structure"),
        ("8-10", "support / compatibility / guard", "Origin geometry, occupancy, rights, terminals and king safety"),
        ("11-13", "proposal / commit / lineage", "Immutable transition, repetition update, pre/post and event hashes"),
        ("14", "projection", "FEN, UCI, reports, GPU files and UI; never authority"),
    ]
    add_table(doc, ["Phases", "Group", "Purpose"], phases, widths=[0.75, 1.8, 4.05], font_size=7.7, first_col_bold=True)
    add_code(doc, "verified_move = origin_support_ok\n             and piece_geometry_ok\n             and occupancy_compatibility_ok\n             and special_rights_ok\n             and king_safety_guard_ok\n             and pre_hash_matches_authority")
    add_callout(doc, "Non-commutative order",
                "Moving before king-safety verification, changing castling rights after repetition hashing, or preserving an en-passant target without legal-capture semantics changes the proof state. Order is data, not an implementation detail.",
                fill="FFF4F2", accent=RED)
    doc.add_page_break()


def perft_validation(doc: Document) -> None:
    add_section_title(doc, "6", "Exact Legal-Move Validation",
                      "Perft checks the successor relation without depending on a heuristic evaluation function.")
    perft_records = json.loads((ROOT / "validation" / "perft_results_v2.json").read_text(encoding="utf-8"))
    grouped: dict[str, list[dict[str, object]]] = {}
    for row in perft_records:
        grouped.setdefault(str(row["fixture"]), []).append(row)
    rows = []
    labels = {
        "initial": "Initial",
        "kiwipete": "Kiwipete",
        "position_3": "Position 3",
        "position_4": "Position 4",
        "position_5": "Position 5",
        "position_6": "Position 6",
    }
    stress = {
        "initial": "Ordinary development, checks and captures",
        "kiwipete": "Castling, pins, checks and dense tactics",
        "position_3": "En-passant and unusual king exposure",
        "position_4": "Castling rights and promotions",
        "position_5": "Promotions, checks and tactical legality",
        "position_6": "Pins, checks and complex piece geometry",
    }
    for key, items in grouped.items():
        items = sorted(items, key=lambda r: int(r["depth"]))
        highest = items[-1]
        rows.append((labels[key], "1-" + str(highest["depth"]), f"{int(highest['expected']):,} / {int(highest['actual']):,}", stress[key]))
    add_table(doc, ["Position", "Depths", "Expected / actual highest", "Coverage stress"], rows,
              widths=[1.05, 0.75, 1.65, 3.15], font_size=7.5, first_col_bold=True)

    add_callout(doc, "Captured result",
                f"{validation['python_perft']['checks']} Python perft checks pass. The independent C++ packed self-test passes nine exact fixtures, including 197,281 initial-position nodes at depth four. CTest reports {validation['native_build']['ctest_total']} tests and zero failures.",
                fill="EAF7F3", accent=TEAL)
    doc.add_heading("Interpretation", level=2)
    add_bullets(doc, [
        "Perft agreement is a mandatory legality regression gate because every later proof depends on the exact successor relation.",
        "Perft does not establish strategy, game value or completeness of the global proof graph.",
        "The Python, C++ and packed protocol paths are deliberately separate enough to expose shared-state and serialization mistakes.",
    ], compact=False)
    add_code(doc, "PYTHONPATH=src python -m ugts_chess perft --depth 4\ncpp/build/host-release/ugts-chess2 selftest\ncpp/build/host-release/ugts-chess-gpu self-test")
    doc.add_page_break()


def wdl_calculus(doc: Document) -> None:
    add_section_title(doc, "7", "Four-Valued WDL Proof Calculus",
                      "WIN, LOSS and DRAW have different quantified obligations. UNKNOWN is the correct result when any obligation remains open.")
    add_pic(doc, "wdl_calculus.png", width=6.75,
            alt="Four-valued WDL calculus showing WIN as one verified loss child, LOSS as complete all-win coverage, DRAW as complete no-win coverage plus a draw argument, and UNKNOWN as an open obligation.")
    caption(doc, "Values are from the side-to-move perspective. A cutoff or score never silently becomes DRAW.")

    rows = [
        ("WIN(q)", "Existential", "At least one verified legal action reaches an exact LOSS child."),
        ("LOSS(q)", "Universal", "Every legal action is covered and every child is exact WIN."),
        ("DRAW(q)", "Complete no-win strategy", "No child is LOSS, every relevant action is resolved, and a draw action, terminal or closed-complement argument is certified."),
        ("UNKNOWN(q)", "Open", "Any required edge, history state, dead-position decision, table dependency or checker result is missing."),
    ]
    add_table(doc, ["Value", "Quantifier", "Certificate obligation"], rows, widths=[1.0, 1.35, 4.25], font_size=8.0, first_col_bold=True)
    add_code(doc, "WIN  <- exists child LOSS\nLOSS <- all legal children WIN\nDRAW <- complete graph and no child LOSS plus draw witness/closed complement\nUNKNOWN <- otherwise")
    add_callout(doc, "Draw discipline",
                "No mate found within N plies, an evaluation of 0.00, a repeated cache key or an unverified fixed point on an incomplete graph is not a draw proof.",
                fill="FFF4F2", accent=RED)
    doc.add_page_break()


def certificate_promotion(doc: Document) -> None:
    add_section_title(doc, "8", "Candidate Certificates and Independent Promotion",
                      "Workers search and propose. A separate checker regenerates the state/action relation before the campaign accepts a value.")
    add_pic(doc, "certificate_flow.png", width=6.75,
            alt="Certificate promotion workflow from worker to candidate, independent checker, campaign ledger and root aggregate, with invalid shortcuts such as score, 64-bit hit, GPU hash and unverified table rejected.")
    caption(doc, "Proof promotion is a verified event with evidence; workers never write authoritative WDL directly.")

    rows = [
        ("State identity", "Full reconstructible state, history counts and exact rule profile."),
        ("Action coverage", "One verified witness for WIN; complete legal successor list for LOSS; complete no-win coverage for DRAW."),
        ("Terminal evidence", "Checkmate, stalemate, dead position or automatic/claimed draw classification replayed from rules."),
        ("Hash closure", "Canonical node hash, child hashes and certificate hash recomputed independently."),
        ("Checker separation", "Candidate and independent-check records have distinct paths and SHA-256 values."),
        ("Campaign event", "Promotion changes campaign status through a hash-chained ledger entry."),
    ]
    add_table(doc, ["Obligation", "Required evidence"], rows, widths=[1.5, 5.1], font_size=7.8, first_col_bold=True)
    add_callout(doc, "Invalid shortcuts",
                "A positive engine score, a principal variation, a 64-bit transposition hit, a GPU digest or an unverified external table value cannot satisfy the quantified certificate obligations.",
                fill="FFF9E8", accent=GOLD)
    doc.add_page_break()


def bounded_examples(doc: Document) -> None:
    add_section_title(doc, "9", "Bounded Exact Results",
                      "The package demonstrates both a closed proof and a useful incomplete search without conflating them.")
    add_pic(doc, "bounded_results.png", width=6.75,
            alt="Side-by-side bounded WDL examples: a KQK position with a 50-move claim available but a checkmate win, and the initial position at depth two returning unknown.")
    caption(doc, "The first search closes an existential proof. The second records explored work and remains UNKNOWN.")

    rows = [
        ("Mate-over-claim fixture", str(mate_wdl["root"]["value"]).upper(), str(mate_wdl["nodes_searched"]), "Qa4# closes WIN although a current 50-move claim action is available."),
        ("Initial position, 2 plies", str(initial_wdl["root"]["value"]).upper(), str(initial_wdl["nodes_searched"]), f"{initial_wdl['cutoffs']} cutoff nodes stay UNKNOWN."),
        ("Retained mate-in-two", "PROVED", "418 explored", "1. Qb5 Ka2 2. Qb2#; four certificate nodes independently verified."),
    ]
    add_table(doc, ["Fixture", "Result", "Work", "Interpretation"], rows, widths=[1.45, 0.85, 0.9, 3.4], font_size=7.7, first_col_bold=True)
    add_callout(doc, "Exact bounded claim",
                "The result is exact only for the serialized root, rule profile and completed certificate. A bounded UNKNOWN does not imply the true game value is draw or that no winning line exists beyond the horizon.",
                fill="EEF4FF", accent=BLUE)
    doc.add_page_break()


def root_campaign(doc: Document) -> None:
    add_section_title(doc, "10", "The Initial Position as Twenty Root Obligations",
                      "Each legal first move is a stable, independently checkable proof shard with complete child state and history context.")
    add_pic(doc, "root_obligations.png", width=6.75,
            alt="Grid of the twenty legal initial moves, each labeled with its root obligation identifier, exact depth-four leaf workload and unknown WDL status.")
    caption(doc, "Root decomposition. The exact depth-four workloads sum to the known 197,281 legal paths.")

    add_code(doc, "root = WIN  if any verified child value is LOSS\nroot = LOSS if all 20 verified child values are WIN\nroot = DRAW if all 20 are resolved, none is LOSS, and a draw strategy closes\nroot = UNKNOWN otherwise")
    add_bullets(doc, [
        "Every shard stores parent/child semantic hashes, child FEN, exact repetition counts, move UCI/SAN, rule profile and a content hash.",
        "Root workload counts are validation evidence and scheduling hints, not WDL values.",
        "The campaign starts with twenty unresolved jobs and zero independently verified child values; therefore the only valid aggregate is UNKNOWN.",
        "A single verified child LOSS would prove White WIN. Proving LOSS or DRAW requires complete root coverage.",
    ], compact=False)
    add_callout(doc, "Current root", f"Root WDL: UNKNOWN\nRoot obligations: {workload['root_obligations']}\nExact depth-four legal leaf paths: {workload['total_exact_leaf_paths']:,}", fill="FFF4F2", accent=RED)
    doc.add_page_break()


def campaign_ledger(doc: Document) -> None:
    add_section_title(doc, "11", "Portable SQLite Campaign Ledger",
                      "The campaign database coordinates work while its event chain and immutable proof files preserve an auditable promotion history.")
    add_pic(doc, "campaign_ledger.png", width=6.75,
            alt="SQLite campaign ledger diagram with jobs, events and metadata tables and a hash-chained sequence of initialization, lease, candidate and verification actions.")
    caption(doc, "Coordination state is mutable; proof files and event hashes make changes inspectable and replayable.")

    rows = [
        ("meta", "Canonical component identity, rule profile, root FEN and root hashes."),
        ("jobs", "Twenty obligations, relative shard/certificate/checker paths, lease state, WDL and verification status."),
        ("events", "Monotonic sequence, action, payload, previous event hash and current event hash."),
        ("relocation", "All stored paths are relative to the database directory; the extracted checkpoint re-verifies on another machine."),
        ("merge policy", "Only checker-accepted values enter the root aggregate; contradictions are rejected or quarantined."),
    ]
    add_table(doc, ["Record", "Contract"], rows, widths=[1.25, 5.35], font_size=8.0, first_col_bold=True)
    add_code(doc, "PYTHONPATH=src python -m ugts_chess campaign-init --db examples/campaign/initial.sqlite3 --shard-dir examples/campaign/root_shards --force\nPYTHONPATH=src python -m ugts_chess campaign-status examples/campaign/initial.sqlite3\nPYTHONPATH=src python -m ugts_chess campaign-verify examples/campaign/initial.sqlite3")
    add_callout(doc, "Captured ledger", "20 unresolved jobs; one initialization event; valid hash chain; zero verified child values; root UNKNOWN.", fill="EAF7F3", accent=TEAL)
    doc.add_page_break()


def tablebases(doc: Document) -> None:
    add_section_title(doc, "12", "Bundled Exact KQK and KRK Tablebases",
                      "The retained three-piece tables are complete finite graphs, not sampled positions or heuristic endgame scores.")
    add_pic(doc, "tablebase_counts.png", width=6.55,
            alt="Bar chart of exact KQK and KRK win, loss, draw and invalid address counts with maximum distance-to-mate values.")
    caption(doc, "Each material class uses all 524,288 cells of an exact 19-bit dense address space; invalid cells remain explicit.")

    rows = [
        ("Valid states", f"{kq['valid_positions']:,}", f"{kr['valid_positions']:,}"),
        ("Win / loss / draw", f"{kq['outcome_counts']['win']:,} / {kq['outcome_counts']['loss']:,} / {kq['outcome_counts']['draw']:,}", f"{kr['outcome_counts']['win']:,} / {kr['outcome_counts']['loss']:,} / {kr['outcome_counts']['draw']:,}"),
        ("Initial mates / stalemates", f"{kq['initial_checkmates']:,} / {kq['initial_stalemates']:,}", f"{kr['initial_checkmates']:,} / {kr['initial_stalemates']:,}"),
        ("Maximum DTM", f"{kq['max_dtm_plies']} plies", f"{kr['max_dtm_plies']} plies"),
        ("Gzip bytes", f"{kq['file_bytes']:,}", f"{kr['file_bytes']:,}"),
        ("Transport SHA-256", kq['sha256'][:18] + "...", kr['sha256'][:18] + "..."),
    ]
    add_table(doc, ["Metric", "KQK", "KRK"], rows, widths=[1.75, 2.4, 2.4], font_size=7.5, first_col_bold=True)
    add_bullets(doc, [
        "Outcome is from the side-to-move perspective; DTM counts plies to checkmate under optimal play.",
        "The 19-bit key is an exact address into a declared profile, not a universal chess-state representation.",
        "Gzip is a transport layer. The raw two-byte-per-cell semantic table remains reconstructible and SHA-checked.",
        "These tables do not replace external larger endgame partitions and do not solve the initial 32-piece position.",
    ], compact=True)
    doc.add_page_break()


def native_foundation(doc: Document) -> None:
    add_section_title(doc, "13", "Independent C++20 Host Foundation",
                      "The native implementation supplies a second legal/perft path, bounded search, mate proof, generic retrograde logic and packed CPU fallback.")
    add_pic(doc, "native_architecture.png", width=6.75,
            alt="Native C++20 architecture showing core legal rules, search, CPU retrograde and SHA-256 modules feeding the ugts-chess2 and ugts-chess-gpu executables.")
    caption(doc, "C++ is an independent legality and performance oracle; the complete claim/history checker remains authoritative in Python.")

    rows = [
        ("ugts-chess2", "info, perft, search, finite mate proof, generic retrograde demo and root-shard enumeration."),
        ("ugts-chess-gpu", "device inspection, 64-byte packed batch expansion, CPU fallback and optional CUDA backend."),
        ("host-release preset", "Ninja, C++20, CUDA disabled, CTest enabled; used for captured host evidence."),
        ("rtx5070ti-release preset", "Ninja, CUDA enabled, CMAKE_CUDA_ARCHITECTURES=120, CTest enabled."),
    ]
    add_table(doc, ["Target", "Role"], rows, widths=[1.8, 4.8], font_size=8.0, first_col_bold=True)
    add_code(doc, "cd cpp\ncmake --preset host-release\ncmake --build --preset host-release --parallel\nctest --preset host-release --output-on-failure")
    add_callout(doc, "Boundary",
                "The C++ host layer intentionally does not claim full tournament-procedure semantics or global WDL authority. It is a separate rules/performance oracle and a basis for the CUDA differential path.",
                fill="FFF9E8", accent=GOLD)
    doc.add_page_break()


def packed_protocol(doc: Document) -> None:
    add_section_title(doc, "14", "Packed Proposal Protocol",
                      "A simple, reconstructible binary interface lets the same batch run through CPU fallback or CUDA before exact comparison with the Python oracle.")
    add_pic(doc, "gpu_protocol.png", width=6.75,
            alt="Packed proposal protocol showing an exact 64-byte position record, a 16-bit move format, Python batch writer, CPU or CUDA expander and independent exact comparison.")
    caption(doc, "The binary batch is a proposal interchange format, not the complete proof certificate.")

    rows = [
        ("PackedPosition", "64 bytes", "Twelve bitboards plus side, castling, en-passant and bounded move-counter fields."),
        ("Move16", "16 bits", "Six source bits, six target bits and promotion encoding."),
        ("Batch header", "UGTSCB20", "Versioned record count and binary framing."),
        ("Move output", "UGMV", "Per-position counts and up to 256 proposal moves in the correctness-first implementation."),
        ("Verifier", "Exact", "Sort/deduplicate and compare with Python legal UCI set; any missing or extra move blocks use."),
    ]
    add_table(doc, ["Record", "Size / tag", "Contract"], rows, widths=[1.35, 1.05, 4.2], font_size=7.8, first_col_bold=True)
    add_callout(doc, "Captured differential",
                f"{validation['packed_differential']['positions']} positions; {validation['packed_differential']['proposed_moves']} proposed moves; {validation['packed_differential']['verified_moves']} verified moves; zero mismatches. The captured backend was the CPU fallback.",
                fill="EAF7F3", accent=TEAL)
    doc.add_page_break()


def rtx_profile(doc: Document) -> None:
    add_section_title(doc, "15", "RTX 5070 Ti Laptop Execution Profile",
                      "The profile is sized for a 12 GB laptop GPU but runtime device inspection and free-memory evidence override every nominal assumption.")
    add_pic(doc, "rtx_memory.png", width=6.75,
            alt="Starting memory plan for a 12 GB RTX 5070 Ti Laptop GPU with 9 GiB solver budget and 3 GiB reserved headroom split among proof index, frontier, moves, retrograde, checkpoint and scratch allocations.")
    caption(doc, "Starting allocation plan. The package reduces batch size rather than weakening proof checks when memory or thermal limits are encountered.")

    rows = [
        ("NVIDIA published laptop specification", "5,888 CUDA cores; 12 GB GDDR7; 672 GB/s memory bandwidth; Blackwell generation."),
        ("Compute profile", "GeForce RTX 5070 Ti is listed under CUDA compute capability 12.0; the checked-in CMake preset requests SM120."),
        ("Nominal allocation", f"{profile['solver_budget_mib'] // 1024} GiB solver budget + {profile['reserved_headroom_mib'] // 1024} GiB reserved headroom."),
        ("Initial launch", f"{profile['batch_positions_initial']:,} positions; minimum {profile['batch_positions_min']:,}; {profile['threads_per_block']} threads/block; {profile['streams_initial']} streams."),
        ("Toolkit profile", "CUDA 12.8 or newer; current CUDA 13.x compatibility requires driver branch 580 or newer."),
        ("Runtime gate", "Confirm exact device name, compute capability, free/total VRAM, driver, toolkit, laptop power mode and thermal state."),
    ]
    add_table(doc, ["Field", "Release contract"], rows, widths=[2.1, 4.5], font_size=7.6, first_col_bold=True)
    add_callout(doc, "No device-performance claim",
                "The packaging host had no nvcc and no physical RTX 5070 Ti Laptop GPU. No CUDA throughput, VRAM peak, temperature, power, battery or laptop speed result is claimed here.",
                fill="FFF4F2", accent=RED)
    doc.add_page_break()


def cuda_boundary(doc: Document) -> None:
    add_section_title(doc, "16", "CUDA Work Division and Proof Boundary",
                      "The GPU performs bounded parallel proposal work; deterministic host-side verification decides what enters the proof graph.")
    rows = [
        ("Move expansion kernel", "One packed position per thread in the correctness-first implementation; produces bounded Move16 proposals."),
        ("Fixed-point kernel", "Iterates monotone candidate states on a supplied finite graph; output is rechecked before promotion."),
        ("Batch sizing", "Starts at 131,072 positions; halves after allocation failure and may reduce streams under thermal pressure."),
        ("Deterministic order", "Kernel output is canonicalized and compared as exact move sets; proof records use sorted content-addressed children."),
        ("Authority", "GPU writes cannot mutate chess state, repetition history, campaign WDL or root aggregate directly."),
        ("Failure state", "No device, insufficient memory, unsupported architecture, mismatch or non-convergence returns an explicit failure/UNKNOWN status."),
    ]
    add_table(doc, ["Component", "Contract"], rows, widths=[1.75, 4.85], font_size=7.8, first_col_bold=True)
    add_code(doc, "GPU: packed positions -> proposal moves / candidate fixed-point labels\nHOST: reconstruct -> legal/history verification -> quantified obligation check\nLEDGER: candidate path + checker path + hashes -> promotion event")
    doc.add_heading("Equal-output measurement first", level=2)
    add_numbered(doc, [
        "Run the Python and C++ exact fixture suite before enabling CUDA.",
        "Differential-test every CUDA move list against the Python oracle on packaged and seeded legal positions.",
        "Measure latency, throughput, VRAM and thermals only after mismatch count is zero.",
        "Do not trade legal, history or certificate checks for occupancy or batch size.",
    ])
    add_callout(doc, "UGTS precision rule",
                "Packed width or device speed is accepted only when the decoder reconstructs every required field and error/collision policy cannot change event ordering or proof obligations.",
                fill="EEF4FF", accent=BLUE)
    doc.add_page_break()


def codex_handoff(doc: Document) -> None:
    add_section_title(doc, "17", "Codex Build, Measure and Promotion Workflow",
                      "The accompanying scripts turn the source tree into a repeatable target-laptop experiment without pretending that one run closes the game.")
    add_pic(doc, "codex_workflow.png", width=6.75,
            alt="Six-step Codex workflow: capture device, build SM120, differential test, measure, extend proof DAG and close independently checked root shards.")
    caption(doc, "Every engineering step has a proof-preserving promotion gate.")

    add_code(doc, "# Windows PowerShell from the package root\npowershell -ExecutionPolicy Bypass -File scripts/build_rtx5070ti.ps1\npowershell -ExecutionPolicy Bypass -File scripts/run_codex_campaign.ps1")
    doc.add_heading("Required device evidence", level=2)
    add_bullets(doc, [
        "Exact GPU name, compute capability, driver, toolkit, total/free VRAM, OS, laptop power mode and clocks.",
        "p50/p95/p99 batch latency, positions/s, moves/s, peak VRAM and host RAM.",
        "5-, 15- and 30-minute clock, temperature, power and throttling traces.",
        "Storage bytes per verified node, checkpoint/replay success and campaign relocation verification.",
        "Differential mismatch count, CUDA fallback reasons and independent certificate-verifier results.",
    ], compact=False)
    add_callout(doc, "Codex objective",
                "Improve proof throughput and close content-addressed obligations. Do not optimize for Elo, a pretty principal variation or a deep but unverifiable search cutoff.",
                fill="FFF9E8", accent=GOLD)
    doc.add_page_break()


def validation_page(doc: Document) -> None:
    add_section_title(doc, "18", "Captured Host Validation",
                      "All exact claims in this report are tied to reproducible package artifacts. Timings remain environment-specific.")
    add_pic(doc, "validation_dashboard.png", width=6.75,
            alt="Validation dashboard showing 92 Python tests, 3 CTest cases, 20 perft checks, 20 root shards, 97 GPU protocol moves with zero mismatches and two hash-valid tablebases.")
    caption(doc, "Host release gates. CUDA compilation and physical-laptop measurement are explicitly pending.")

    rows = [
        ("Python compile/tests", f"PASS - {validation['python_tests']['count']} tests", f"{validation['python_tests']['seconds']:.2f} s captured suite"),
        ("Python perft", f"PASS - {validation['python_perft']['checks']} checks", f"largest {validation['python_perft']['largest_nodes']:,} nodes"),
        ("Native build/CTest", f"PASS - {validation['native_build']['ctest_total']}/3", "clean CMake/Ninja CPU build"),
        ("Packed differential", "PASS - 0 mismatches", f"{validation['packed_differential']['verified_moves']} exact moves across {validation['packed_differential']['positions']} positions"),
        ("Schemas/proofs", "PASS", f"{validation['schemas_and_proofs']['schema_documents']} instances / {validation['schemas_and_proofs']['schema_files']} schemas; mate and WDL verified"),
        ("Campaign", "PASS", "20 obligations; valid event chain; root UNKNOWN"),
        ("Tablebases", "PASS", f"2 files; {validation['tablebases']['compressed_bytes']:,} compressed bytes; DTM 20/32"),
        ("Wheel", "PASS", f"{validation['wheel']['file']}; clean install"),
        ("CUDA device", "DEFERRED", validation['cuda_device_validation']['status']),
    ]
    add_table(doc, ["Gate", "Status", "Evidence"], rows, widths=[1.35, 1.45, 3.8], font_size=7.3, first_col_bold=True)
    add_callout(doc, "Scientific boundary",
                "Passing tests establish internal consistency and exact bounded outputs under the declared profiles. They do not establish the game-theoretic value of the initial position or a physical-GPU advantage.",
                fill="FFF4F2", accent=RED)
    doc.add_page_break()


def package_and_catalog(doc: Document) -> None:
    add_section_title(doc, "19", "Package Structure and Mechanism Catalog",
                      "The ZIP is a complete Codex handoff with source, exact data, schemas, scripts, examples, validation and editable report material.")
    add_pic(doc, "package_map.png", width=6.75,
            alt="Package map showing Python source, C++ and CUDA, campaign examples, exact data, formal specs, scripts, validation and report/doc folders.")
    caption(doc, "The source PDFs are referenced by hash and are not redistributed.")

    counts = Counter(str(m["domain"]) for m in mechanisms)
    top = counts.most_common(10)
    rows = [(domain, str(count), "implemented/specification entries") for domain, count in top]
    add_table(doc, ["Catalog domain", "Rows", "Role"], rows, widths=[2.1, 0.8, 3.7], font_size=7.5, first_col_bold=True)
    p = doc.add_paragraph()
    p.add_run("Catalog continuity. ").bold = True
    p.add_run(f"The chess component catalog contains {len(mechanisms)} contiguous records C001-C104. Each row records domain, normalized definition, disposition and validation state. Catalog IDs are traceability records, not novelty or ownership claims.")
    add_callout(doc, "Distribution",
                f"The clean-install wheel {validation['wheel']['file']} is {validation['wheel']['bytes']:,} bytes. Runtime Python code is dependency-free; report generation and validation tooling are separate development concerns.",
                fill="EEF4FF", accent=BLUE)
    doc.add_page_break()


def roadmap_and_kills(doc: Document) -> None:
    add_section_title(doc, "20", "Promotion Roadmap and Kill Criteria",
                      "Progress is measured by independently closed obligations, not by search depth, Elo, node count or GPU utilization alone.")
    add_pic(doc, "roadmap.png", width=6.75,
            alt="Roadmap from frozen authority and proof infrastructure through exact endgame growth, complete draw logic, root-shard closure and a final initial-position certificate.")
    caption(doc, "The final promotion gate remains open.")

    doc.add_heading("Hard kill criteria", level=2)
    add_bullets(doc, [
        "Any legal-move mismatch in perft, special-rule fixtures or Python/C++/CUDA differential checks.",
        "Any certificate accepted with a missing defender reply, incomplete root coverage, bad pre/post hash or wrong history profile.",
        "Any finite cutoff, heuristic score or unresolved SCC silently coerced to DRAW.",
        "Any packed/cache key treated as complete identity without reconstructibility and collision policy.",
        "Any external tablebase result used without explicit material/rule profile and independent adapter verification.",
        "Any state or WDL mutation performed directly by a GPU worker, UI, stale lease or unchecked model output.",
        "Any target-laptop performance claim made without the required device, thermal and equal-output evidence.",
    ], compact=True)
    add_callout(doc, "Next promotion gates",
                "Complete dead-position certificates; verified adapters for larger exact endgame partitions; disk-backed content-addressed proof DAG; long-run CUDA differential evidence; independently verified root-shard closure; complete initial WIN/DRAW/LOSS certificate.",
                fill="FFF9E8", accent=GOLD)
    doc.add_page_break()


def formal_definition(doc: Document) -> None:
    add_section_title(doc, "21", "Final Formal Definition and Release Decision",
                      "The solved object is a typed proof state, not a rendered board, engine evaluation or GPU workload.")
    add_callout(doc, "Formal definition",
                "UGTS-KC Chess 2.0 is a content-addressed classical-chess proof profile in which complete state and history define a deterministic legal action relation; moves and draw claims pass through support, compatibility, terminal and king-safety guards; verified proposals commit immutable transitions with repetition and hash lineage; finite closed graphs receive exact WDL/DTM values; bounded searches preserve UNKNOWN; and only independent, complete, hash-checked certificates may promote a root to WIN, LOSS or DRAW.",
                fill="EAF7F3", accent=TEAL)
    add_code(doc, "UGTS-CHESS-2 = (Q, A, M, G, T, H, V, C, P, L, X)\nQ  complete typed positions and rule profile\nA  attack/support/occupancy/king-safety relations\nM  deterministic legal move relation\nG  terminal and draw-claim guards\nT  immutable transitions\nH  exact repetition and move-count history\nV  WIN | LOSS | DRAW | UNKNOWN\nC  certificates and checker records\nP  partition/campaign ledger\nL  event, replay and proof lineage\nX  optional CPU/CUDA proposal adapters")

    rows = [
        ("GO", "The exact state/action foundation, proof semantics, campaign ledger and host/CUDA handoff are coherent and executable."),
        ("SOLVED", "Declared bounded mate fixtures and the bundled complete KQK/KRK material domains."),
        ("BOUNDED", "Finite-horizon WDL searches and selected dead-position certificates."),
        ("UNRESOLVED", "The game-theoretic value of the orthodox 32-piece initial position."),
        ("NEXT GATE", "Verified larger endgame partitions, complete draw/dead-position logic and an all-material content-addressed proof DAG."),
    ]
    add_table(doc, ["Decision", "Meaning"], rows, widths=[1.2, 5.4], font_size=8.0, first_col_bold=True)
    add_callout(doc, "Final statement",
                "Classical chess has been converted into an explicit UGTS proof campaign and solved exactly where the release closes the full declared graph. No unsupported whole-game result is substituted for the remaining proof work.",
                fill="EEF4FF", accent=BLUE)
    doc.add_page_break()


def appendix_sources(doc: Document) -> None:
    add_section_title(doc, "Appendix A", "Source Register and External References",
                      "Supplied UGTS artifacts are identified by filename and SHA-256; external standards provide current chess and device context.")
    rows = []
    for s in source_reg["sources"]:
        rows.append((s["title"], s["page_basis"], s["sha256"][:14] + "...", s["role"]))
    add_table(doc, ["Source", "Page basis", "SHA-256", "Release use"], rows,
              widths=[1.25, 0.85, 1.15, 3.35], font_size=6.5, first_col_bold=True)

    doc.add_heading("External normative and device context", level=2)
    ext_rows = [
        ("FIDE Laws of Chess", "https://handbook.fide.com/chapter/e012023", "Move, terminal, repetition, claim and move-count semantics."),
        ("NVIDIA RTX 50 Series Laptop Specs", "https://www.nvidia.com/en-us/geforce/laptops/50-series/", "RTX 5070 Ti Laptop core, memory and bandwidth data."),
        ("NVIDIA CUDA GPU Compute Capability", "https://developer.nvidia.com/cuda/gpus", "GeForce RTX 5070 Ti listed under compute capability 12.0."),
        ("CUDA Compatibility", "https://docs.nvidia.com/deploy/cuda-compatibility/minor-version-compatibility.html", "CUDA 13.x minimum driver branch 580; current compatibility boundary."),
        ("CUDA Toolkit Release Notes", "https://docs.nvidia.com/cuda/cuda-toolkit-release-notes/", "Windows driver-separation and current toolkit/driver notes."),
    ]
    add_table(doc, ["Reference", "Location", "Use"], ext_rows, widths=[1.55, 2.65, 2.4], font_size=6.4, first_col_bold=True)

    doc.add_heading("Attribution and legal boundary", level=2)
    p = doc.add_paragraph()
    p.add_run("Prepared for Tom Klootwijk. ").bold = True
    p.add_run("The requester-supplied name and project attribution are recorded for continuity and are not independently verified. This report and package are technical design artifacts, not legal proof of identity, authorship, ownership, patentability, priority, standard conformance, exclusive rights or chain of title.")


def main() -> None:
    REPORT.mkdir(parents=True, exist_ok=True)
    doc = setup_document_v2()
    cover(doc)
    executive(doc)
    source_grounding(doc)
    solve_meanings(doc)
    state_profile(doc)
    draw_semantics(doc)
    move_authority(doc)
    perft_validation(doc)
    wdl_calculus(doc)
    certificate_promotion(doc)
    bounded_examples(doc)
    root_campaign(doc)
    campaign_ledger(doc)
    tablebases(doc)
    native_foundation(doc)
    packed_protocol(doc)
    rtx_profile(doc)
    cuda_boundary(doc)
    codex_handoff(doc)
    validation_page(doc)
    package_and_catalog(doc)
    roadmap_and_kills(doc)
    formal_definition(doc)
    appendix_sources(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
