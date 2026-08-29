from __future__ import annotations

import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
ASSETS = REPORT / "assets"
OUT = REPORT / "UGTS_KC_CHESS_2_0_Game_Theoretic_Attempt_RTX5070TI.docx"

NAVY = "0B1736"
INK = "14213D"
BLUE = "4568DC"
CYAN = "1FB7C9"
TEAL = "2AA876"
GOLD = "E0A819"
ORANGE = "E57A44"
MAGENTA = "C04C93"
RED = "C94C4C"
PURPLE = "7357C5"
GREEN = "3F9B62"
LIGHT = "F5F8FC"
MID = "D9E2EF"
GRAY = "65758B"
DARKGRAY = "34465F"
WHITE = "FFFFFF"
DARK = "081124"

validation = json.loads((ROOT / "validation" / "summary_v2.json").read_text(encoding="utf-8"))
kq = json.loads((ROOT / "data" / "kqk.tb.json").read_text(encoding="utf-8"))
kr = json.loads((ROOT / "data" / "krk.tb.json").read_text(encoding="utf-8"))
source_reg = json.loads((ROOT / "spec" / "source_register.json").read_text(encoding="utf-8"))
mechanisms = json.loads((ROOT / "spec" / "chess_mechanisms.json").read_text(encoding="utf-8"))
workloads = json.loads((ROOT / "examples" / "campaign" / "initial_depth4_workloads.json").read_text(encoding="utf-8"))
rtx_profile = json.loads((ROOT / "spec" / "rtx5070ti_profile.json").read_text(encoding="utf-8"))
native_depth5 = json.loads((ROOT / "validation" / "native_perft_depth5.json").read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trPr.append(node)


def set_cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def set_paragraph_bottom_border(paragraph, color=NAVY, size="8") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def add_pic(doc: Document, filename: str, width=6.75, alt=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    shape = run.add_picture(str(ASSETS / filename), width=Inches(width))
    if alt:
        shape._inline.docPr.set("descr", alt)
        shape._inline.docPr.set("title", filename)
    return p


fig_no = 0


def caption(doc: Document, text: str):
    global fig_no
    fig_no += 1
    p = doc.add_paragraph(style="Figure Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Figure {fig_no}. ")
    r.bold = True
    p.add_run(text)
    return p


def table_caption(doc: Document, text: str, number: int):
    p = doc.add_paragraph(style="Table Caption")
    r = p.add_run(f"Table {number}. ")
    r.bold = True
    p.add_run(text)
    return p


def add_callout(doc: Document, title: str, body: str, fill="EAF7F3", accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "10", "color": accent},
        bottom={"val": "single", "sz": "10", "color": accent},
        left={"val": "single", "sz": "18", "color": accent},
        right={"val": "single", "sz": "10", "color": accent},
    )
    set_cell_margins(cell, top=90, start=125, bottom=90, end=125)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9.2)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = cell.add_paragraph(body)
    p2.style = doc.styles["Normal"]
    p2.paragraph_format.space_after = Pt(0)
    return table


def add_table(doc: Document, headers, rows, widths=None, header_fill=NAVY, font_size=8.0, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_cant_split(hdr)
    for j, header in enumerate(headers):
        cell = hdr.cells[j]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, top=55, start=65, bottom=55, end=65)
        p = cell.paragraphs[0]
        r = p.add_run(str(header))
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(font_size)
        if widths:
            cell.width = Inches(widths[j])
    border = {"val": "single", "sz": "4", "color": MID}
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        set_cant_split(table.rows[-1])
        fill = WHITE if i % 2 == 0 else LIGHT
        for j, value in enumerate(row):
            cell = cells[j]
            set_cell_shading(cell, fill)
            set_cell_margins(cell, top=50, start=65, bottom=50, end=65)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor.from_string(INK)
            if first_col_bold and j == 0:
                r.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[j])
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    return table


def add_bullets(doc: Document, items, compact=True, accent=TEAL):
    for item in items:
        p = doc.add_paragraph(style="Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(1 if compact else 3)
        marker = p.add_run("• ")
        marker.bold = True
        marker.font.color.rgb = RGBColor.from_string(accent)
        p.add_run(item)


def add_numbered(doc: Document, items):
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(item)


def add_code(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    set_cell_shading(cell, "111C35")
    edge = {"val": "single", "sz": "6", "color": "111C35"}
    set_cell_border(cell, top=edge, bottom=edge, left=edge, right=edge)
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Liberation Mono"
    r.font.size = Pt(7.8)
    r.font.color.rgb = RGBColor.from_string("EAF2FF")


def add_section_title(doc: Document, number: str, title: str, subtitle: str | None = None):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(f"{number}. {title}")
    if subtitle:
        doc.add_paragraph(subtitle, style="Section Lead")


def add_source_note(doc: Document, text: str):
    p = doc.add_paragraph(style="Source Note")
    r = p.add_run("Grounding: ")
    r.bold = True
    p.add_run(text)


def add_page_break(doc: Document):
    doc.add_page_break()


def make_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(3.5)
    normal.paragraph_format.line_spacing = 1.03

    definitions = [
        ("Title", 27, NAVY, 0, 7),
        ("Subtitle", 14.5, BLUE, 0, 7),
        ("Heading 1", 19.5, NAVY, 7, 5),
        ("Heading 2", 13.2, TEAL, 6, 3),
        ("Heading 3", 11.3, BLUE, 5, 2),
    ]
    for name, size, color, before, after in definitions:
        style = styles[name]
        style.font.name = "Noto Sans"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    custom = [
        ("Cover Kicker", WD_STYLE_TYPE.PARAGRAPH, 9.2, GRAY, True),
        ("Cover Meta", WD_STYLE_TYPE.PARAGRAPH, 8.9, DARKGRAY, False),
        ("Section Lead", WD_STYLE_TYPE.PARAGRAPH, 10.2, GRAY, False),
        ("Figure Caption", WD_STYLE_TYPE.PARAGRAPH, 7.9, GRAY, False),
        ("Table Caption", WD_STYLE_TYPE.PARAGRAPH, 7.9, GRAY, False),
        ("Bullet", WD_STYLE_TYPE.PARAGRAPH, 9.0, INK, False),
        ("Source Note", WD_STYLE_TYPE.PARAGRAPH, 7.9, GRAY, False),
        ("Small", WD_STYLE_TYPE.PARAGRAPH, 7.8, GRAY, False),
    ]
    for name, kind, size, color, bold in custom:
        style = styles[name] if name in styles else styles.add_style(name, kind)
        style.font.name = "Noto Sans"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        if kind == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.space_after = Pt(2)
            style.paragraph_format.line_spacing = 1.0
            if name in ("Figure Caption", "Table Caption"):
                style.font.italic = True
    styles["Section Lead"].font.italic = True


def setup_document() -> Document:
    doc = Document()
    make_styles(doc)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(6)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = "UGTS CHESS PROOF CAMPAIGN 2.0  |  CLASSICAL GAME-THEORETIC ATTEMPT"
    p.style = doc.styles["Small"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_bottom_border(p, color=MID, size="6")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Prepared for Tom Klootwijk  |  Page ")
    r.font.size = Pt(7.8)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(p)

    props = doc.core_properties
    props.title = "UGTS Chess Game-Theoretic Solver 2.0 - RTX 5070 Ti Laptop Edition"
    props.subject = "Classical chess proof campaign, exact WDL certificates, disk-backed orchestration and optional CUDA SM120 adapters"
    props.author = "OpenAI, prepared for Tom Klootwijk"
    props.keywords = "UGTS, chess, game theory, proof certificate, WDL, CUDA, RTX 5070 Ti, SM120, tablebase, campaign"
    props.comments = "The initial classical position remains UNKNOWN. Requester attribution is recorded as supplied and not independently verified."
    return doc


def cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Cover Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("UGTS APPLICATION / PROOF CAMPAIGN")
    add_pic(doc, "cover_architecture.png", width=6.8, alt="Architecture of the UGTS Chess 2.0 proof campaign and RTX execution profile")
    rows = [
        ("Canonical identity", "ugts.application.chess-proof@2.0.0"),
        ("Release", "UGTS Chess Game-Theoretic Solver 2.0"),
        ("Target", "NVIDIA GeForce RTX 5070 Ti Laptop GPU, nominal 12 GB GDDR7"),
        ("Captured root value", "UNKNOWN"),
        ("Document date", "29 August 2026"),
        ("Status", "Executable host-validated foundation; CUDA laptop build and long campaign delegated to Codex/device"),
    ]
    add_table(doc, ["Release field", "Value"], rows, widths=[1.55, 5.05], font_size=8.1, first_col_bold=True)
    add_callout(
        doc,
        "Evidence boundary",
        "This release attempts a proof-oriented solution of classical chess. It does not claim that the standard initial position is weakly, strongly or ultra-weakly solved. GPU kernels propose work; independently checked certificates decide truth.",
        fill="FFF4F2",
        accent=RED,
    )
    add_page_break(doc)


def executive(doc: Document) -> None:
    add_section_title(doc, "Executive Summary", "Release Decision", "What has been built, what has been proved, and what remains open.")
    add_callout(
        doc,
        "Formal decision - go",
        "Promote Chess 2.0 as an upgraded foundational proof campaign and RTX 5070 Ti Laptop handoff. Freeze the exact rule/history identity, represent claimable draws as actions, issue content-addressed WDL obligations, and keep the classical initial position UNKNOWN until quantified proof obligations close.",
        fill="EAF8F2",
        accent=TEAL,
    )
    doc.add_heading("Verified release result", level=2)
    metrics = validation
    rows = [
        ("Initial position", "20 exact root obligations; 0 independently verified child values; root UNKNOWN"),
        ("Python", f"{metrics['python_tests']['count']} tests passed; {metrics['python_perft']['checks']} exact perft checks"),
        ("Native C++20", f"{metrics['native_build']['ctest_total']} CTest gates passed; depth-5 initial perft {native_depth5['nodes']:,} nodes"),
        ("Packed protocol", f"{metrics['packed_differential']['positions']} positions; {metrics['packed_differential']['proposed_moves']} moves; 0 differential mismatches"),
        ("Certificates", "Mate certificate verified; exact bounded WIN fixture; initial-depth-2 fixture safely UNKNOWN"),
        ("Tablebases", f"Complete KQK max DTM {metrics['tablebases']['kqk_max_dtm']} plies; KRK max DTM {metrics['tablebases']['krk_max_dtm']} plies"),
        ("Distribution", f"Wheel {metrics['wheel']['bytes']:,} bytes; isolated clean install passed"),
        ("CUDA laptop", "Source and SM120 preset supplied; no physical laptop benchmark claimed in this environment"),
    ]
    add_table(doc, ["Area", "Captured result"], rows, widths=[1.45, 5.15], font_size=8.1, first_col_bold=True)
    doc.add_heading("What the deliverable is for", level=2)
    add_bullets(doc, [
        "Codex can compile the optional CUDA adapters, compare them against the exact Python and C++ host oracles, and record device evidence without weakening proof semantics.",
        "The twenty first moves can be leased as independent shards, with candidate WDL results kept non-authoritative until a separate checker record is attached.",
        "The package supplies a reproducible starting point for out-of-core proof DAGs, larger exact endgame partitions, dead-position certificates and distributed replay.",
    ])
    add_page_break(doc)


def contents(doc: Document) -> None:
    add_section_title(doc, "Reading Map", "Contents and Claim Ladder", "The report moves from proof meaning to exact state, certificates, campaign orchestration, RTX execution and promotion gates.")
    rows = [
        ("1-3", "Meaning, source basis and component identity"),
        ("4-7", "FIDE action profile, UGTS authority and exact history state"),
        ("8-12", "WDL calculus, certificates, root obligations and SQLite campaign"),
        ("13-15", "Bounded exact results, retained mate proof and exact KQK/KRK tables"),
        ("16-20", "C++20 core, packed protocol, RTX profile, VRAM plan and CUDA boundary"),
        ("21-24", "Codex workflow, global proof architecture, validation and performance evidence"),
        ("25-28", "Roadmap, mechanism catalog, package map, final definition and references"),
    ]
    add_table(doc, ["Sections", "Purpose"], rows, widths=[0.85, 5.75], font_size=8.4, first_col_bold=True)
    add_pic(doc, "solve_ladder.png", width=6.65, alt="Three distinct meanings of solved chess")
    caption(doc, "Rules solved, positions solved and the initial game solved are separate claims. Chess 2.0 establishes the first and selected instances of the second; the third remains UNKNOWN.")
    add_callout(doc, "Reading rule", "An exact result is promoted only when its serialized rule profile, state identity, action coverage, certificate graph and independent verification status support it. Search depth, evaluation score and GPU completion do not substitute for those obligations.", fill="EEF4FF", accent=BLUE)
    add_page_break(doc)


def source_basis(doc: Document) -> None:
    add_section_title(doc, "1", "Source Basis and Component Identity", "UGTS supplies authority discipline; chess rules, WDL logic and GPU engineering are explicit application additions.")
    add_pic(doc, "source_grounding.png", width=6.72, alt="UGTS source layers and Chess 2.0 engineering delta")
    caption(doc, "The new chess layer retains the recurring support -> compatibility -> guard -> verified event -> transition -> lineage architecture.")
    doc.add_heading("1.1 Immediate source use", level=2)
    rows = [
        ("UGTS-GN 1.1", "Authority chain, packed precision and hardware evidence boundaries"),
        ("UGTS-KC 2.0", "Typed state, topology, bounded queries and explicit degeneracy"),
        ("Two Hands 3.0", "Proposal/commit, pre/post hashes, checkpoints and replay"),
        ("Literal 3.6 / SCLP 3.6.2", "Content addresses, dependency order, finite keys and bounded branching"),
        ("Elizabeth 3.9", "Deterministic game world, snapshots, CLI and offline delivery"),
        ("Go 4.2 / Operator 4.2", "Rules-position-game claim ladder; exact certificates; distinct evaluation, event and replay orders"),
        ("Chess 1.0", "Exact legal kernel, mate proofs, search boundary and complete KQK/KRK tables"),
    ]
    add_table(doc, ["Source", "Use in Chess 2.0"], rows, widths=[1.65, 4.95], font_size=7.8, first_col_bold=True)
    add_source_note(doc, "The source PDFs do not claim a classical chess solution or RTX implementation. The chess/WDL/campaign/CUDA layers are visibly engineering-derived. SARA 3.6.3 was reviewed but its wallet-specific mechanisms were excluded.")
    add_page_break(doc)


def solution_meaning(doc: Document) -> None:
    add_section_title(doc, "2", "What a Game-Theoretic Solution Requires", "A finite game can be defined exactly while its initial value remains computationally unresolved.")
    doc.add_heading("2.1 Three independent claims", level=2)
    rows = [
        ("Rules solved", "Every legal action, terminal and transition is deterministic and testable.", "Implemented"),
        ("Position solved", "A fully serialized finite state has a complete proof certificate or exact table entry.", "Implemented for bounded fixtures, mate certificates, KQK and KRK"),
        ("Game solved", "The orthodox initial position is certified WIN, DRAW or LOSS under one complete rule profile.", "Not established; UNKNOWN"),
    ]
    add_table(doc, ["Claim", "Required evidence", "Chess 2.0 status"], rows, widths=[1.15, 3.4, 2.05], font_size=7.8, first_col_bold=True)
    doc.add_heading("2.2 Weak, strong and ultra-weak terminology", level=2)
    add_bullets(doc, [
        "An ultra-weak solution determines the game-theoretic value of the initial position.",
        "A weak solution also gives a strategy from the initial position that achieves that value against every reply.",
        "A strong solution gives the game-theoretic value and optimal action for every legal position in the selected state space.",
        "Chess 2.0 claims none of these for unrestricted classical chess. Its root node is deliberately serialized as UNKNOWN.",
    ], compact=False)
    add_callout(doc, "Why UNKNOWN is a result", "UNKNOWN means one or more quantified obligations remain open. It is not a heuristic draw, an engine disagreement, a timeout disguised as a result, or evidence that no solution exists.", fill="F4F0FF", accent=PURPLE)
    add_page_break(doc)


def fide_profile(doc: Document) -> None:
    add_section_title(doc, "3", "Classical FIDE Rule and Action Profile", "Moves, claimable draws and automatic terminals are distinct action classes.")
    rows = [
        ("Ordinary move", "Legal board transition", "Must pass occupancy, motion, special-right and king-safety checks"),
        ("Current draw claim", "Optional player action", "Threefold repetition or 50-move condition already present"),
        ("Claim by intended move", "Optional player action", "Declared legal move would create third repetition or complete 50 moves"),
        ("Checkmate", "Automatic terminal", "Side to move has no legal move and is in check; winner is the opponent"),
        ("Stalemate", "Automatic terminal", "No legal move and not in check"),
        ("Fivefold repetition", "Automatic terminal", "Exact repetition identity occurs at least five times"),
        ("75-move rule", "Automatic terminal", "150 halfmoves without pawn move or capture; checkmate takes precedence"),
        ("Dead position", "Automatic terminal", "Requires an exact certificate that no mating sequence is possible"),
    ]
    add_table(doc, ["Action/terminal", "Class", "Proof treatment"], rows, widths=[1.55, 1.45, 3.6], font_size=7.45, first_col_bold=True)
    doc.add_heading("3.1 Declared profile", level=2)
    add_code(doc, "rules_profile = fide-classical-2023-claims-as-actions-v2\nmove claims are optional actions\nautomatic draws are terminals\ncheckmate is evaluated before the automatic 75-move terminal")
    add_callout(doc, "Dead-position boundary", "The bundled recognizer certifies a conservative exact subset. A global classical proof must add a complete dead-position oracle or attach independently checkable dead-position certificates. Uncovered cases remain UNKNOWN.", fill="FFF4F2", accent=RED)
    add_source_note(doc, "FIDE Laws of Chess, Articles 5.2, 9.2, 9.3 and 9.6. The official 2023 text states that checkmate takes precedence over the automatic 75-move draw.")
    add_page_break(doc)


def authority(doc: Document) -> None:
    add_section_title(doc, "4", "One Action Through the UGTS Authority Chain", "A move generator proposes; only the verifier and atomic transition alter authoritative state.")
    add_pic(doc, "authority_chain.png", width=6.72, alt="Chess action mapped to WHERE MATCH READY CHECKED CHANGE REMEMBER")
    caption(doc, "Chess mechanics mapped onto the canonical six-stage authority sequence.")
    rows = [
        ("WHERE / support", "Rays, jump targets, pawn attack/move supports, king neighborhood"),
        ("MATCH / compatibility", "Correct side, piece type, target occupancy and capture relationship"),
        ("READY / special guards", "Promotion choice, castling rights/path, legal en-passant right"),
        ("CHECKED / verified event", "Move leaves own king safe; terminal classification follows the new state"),
        ("CHANGE / transition", "One immutable patch updates board, side, rights, en-passant and clocks"),
        ("REMEMBER / lineage", "Exact repetition counts, state hashes, move event and replay order"),
    ]
    add_table(doc, ["Stage", "Chess interpretation"], rows, widths=[1.55, 5.05], font_size=8.0, first_col_bold=True)
    add_callout(doc, "Non-bypass rule", "CPU search, CUDA kernels, a human interface, opening book, tablebase adapter or neural policy may order or propose actions. None may set an authoritative WDL value without the same legal and certificate checks.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def state_identity(doc: Document) -> None:
    add_section_title(doc, "5", "Exact Position and History Identity", "A board diagram is not enough to decide repetition, move-count claims or proof replay.")
    add_pic(doc, "state_identity.png", width=6.72, alt="Position hash and game state hash fields")
    caption(doc, "The proof identity includes normalized rule state and exact repetition-count context.")
    doc.add_heading("5.1 Authoritative state", level=2)
    add_code(doc, "Q = (board, side, castling rights, legal en-passant right,\n     halfmove clock, repetition-count context, rule profile, lineage)")
    rows = [
        ("Position SHA-256", "Exact FEN-level transport identity including fullmove metadata"),
        ("Game-state SHA-256", "Rule-semantic position plus capped halfmove clock, sorted exact repetition counts and profile ID"),
        ("64-bit cache key", "Permitted only as a fast index; full record must be compared before proof use"),
        ("Lineage hash", "Ordered event evidence; not interchangeable with current-state identity"),
    ]
    add_table(doc, ["Identity", "Role"], rows, widths=[1.7, 4.9], font_size=8.0, first_col_bold=True)
    add_source_note(doc, "This preserves the recurrent UGTS distinction between coordinates, identity and lineage. SCLP finite keys are indexes, not complete worlds.")
    add_page_break(doc)


def draw_history(doc: Document) -> None:
    add_section_title(doc, "6", "Repetition, Move Counts and Draw Strategy", "A claim opportunity changes available actions but need not end a winning line.")
    doc.add_heading("6.1 Exact action semantics", level=2)
    rows = [
        ("Threefold now", "claim_threefold_current", "Player may take DRAW or continue"),
        ("50 moves now", "claim_fifty_move_current", "Player may take DRAW or continue"),
        ("Third repetition by move", "claim_threefold_by_move", "Claim is attached to the declared legal move"),
        ("50 moves by move", "claim_fifty_move_by_move", "Claim is attached to the declared legal move"),
        ("Fivefold", "fivefold_repetition", "Forced DRAW terminal"),
        ("75 moves", "seventy_five_move", "Forced DRAW terminal after checkmate test"),
    ]
    add_table(doc, ["Condition", "Record code", "Game-theoretic treatment"], rows, widths=[1.55, 2.25, 2.8], font_size=7.8, first_col_bold=True)
    doc.add_heading("6.2 Repetition identity", level=2)
    add_bullets(doc, [
        "Board occupancy, side to move and castling rights must match.",
        "An en-passant square contributes only when an en-passant capture is actually legal; a merely printed target is normalized away.",
        "Exact occurrence counts are stored as sorted SHA-256 identity/count pairs. A rolling checksum alone is insufficient.",
        "Fullmove number is replay metadata and does not change legal action or terminal rules; the semantic proof hash excludes it.",
    ], compact=False)
    add_callout(doc, "Strategic consequence", "At halfmove 100, a player may already claim a draw. If the same position also contains checkmate in one, WIN remains exact because the player can choose the winning move instead of the optional draw action.", fill="EAF8F2", accent=TEAL)
    add_page_break(doc)


def wdl(doc: Document) -> None:
    add_section_title(doc, "7", "Four-Valued WDL Proof Calculus", "WIN, LOSS and DRAW are proof claims; UNKNOWN is the safe incomplete value.")
    add_pic(doc, "wdl_calculus.png", width=6.72, alt="Four-valued WDL proof rules")
    caption(doc, "Values are from the side-to-move perspective. Child values are inverted when interpreted by the parent.")
    doc.add_heading("7.1 Quantified obligations", level=2)
    add_code(doc, "WIN(q)  := exists legal action a with exact LOSS(child(q,a))\nLOSS(q) := every legal action is covered and exact WIN(child(q,a))\nDRAW(q) := no winning action exists, coverage is complete, and a draw terminal/action/closed-complement certificate exists\nUNKNOWN := any required edge, history state, dead-position decision or checker obligation is open")
    add_callout(doc, "Cutoff law", "A horizon, node budget, time limit, missing child or heuristic score is UNKNOWN. The solver never promotes 'no win found yet' to DRAW.", fill="FFF4F2", accent=RED)
    add_source_note(doc, "The fixed-point structure follows the bounded exact-solving precedent in UGTS-KC 4.2 Go and the order/certification discipline in the General Operator and Order Addendum.")
    add_page_break(doc)


def certificates(doc: Document) -> None:
    add_section_title(doc, "8", "Content-Addressed Certificates and Independent Promotion", "A worker result is evidence only after a separate checker validates quantified coverage.")
    add_pic(doc, "certificate_flow.png", width=6.72, alt="Worker candidate checker campaign promotion flow")
    caption(doc, "A certificate path and a checker record are both bound into campaign lineage.")
    rows = [
        ("Node identity", "Game-state SHA-256, FEN, history counts, remaining depth and rule profile"),
        ("Child obligation", "Action kind, UCI/SAN or claim code, child state hash, child value and child certificate hash"),
        ("Coverage", "terminal, witness, complete or cutoff"),
        ("Certificate hash", "SHA-256 over canonical node content excluding the hash field itself"),
        ("Checker record", "Obligation ID, claimed WDL, certificate SHA-256, child game-state SHA-256 and valid=true"),
        ("Promotion", "Campaign job changes to verified only after checker fields and hashes match"),
    ]
    add_table(doc, ["Record", "Required content"], rows, widths=[1.45, 5.15], font_size=7.8, first_col_bold=True)
    add_callout(doc, "Independence boundary", "The current reference checker runs in the same Python package but through a separate verification path and record. A production campaign should additionally diversify implementations, hosts and reviewers before treating a root certificate as durable evidence.", fill="F4F0FF", accent=PURPLE)
    add_page_break(doc)


def root_obligations(doc: Document) -> None:
    add_section_title(doc, "9", "The Initial Position Becomes Twenty Root Obligations", "The first exact decomposition is small, stable and independently leaseable.")
    add_pic(doc, "root_obligations.png", width=6.72, alt="Twenty classical initial chess moves as root proof obligations")
    caption(doc, "Every legal first move has a stable obligation ID, child state hashes and an exact depth-4 workload count.")
    sorted_work = sorted(workloads["shards"], key=lambda item: item["exact_depth4_leaf_paths"], reverse=True)
    rows = [(item["obligation_id"], item["move_uci"], f"{item['exact_depth4_leaf_paths']:,}", item["wdl"].upper()) for item in sorted_work[:8]]
    add_table(doc, ["Largest shard", "Move", "Depth-4 leaf paths", "Captured WDL"], rows, widths=[2.1, 0.8, 1.45, 2.25], font_size=7.7, first_col_bold=True)
    add_callout(doc, "Exact first checkpoint", f"The twenty shard workloads sum to {workloads['total_exact_leaf_paths']:,}, exactly matching standard initial-position perft depth 4. This checks the decomposition; it does not determine any child's full-game value.", fill="EEF4FF", accent=BLUE)
    add_page_break(doc)


def root_aggregation(doc: Document) -> None:
    add_section_title(doc, "10", "Root Aggregation and Current Campaign Status", "Only independently verified child values participate in the initial-position result.")
    rows = [
        ("White root WIN", "At least one verified child is LOSS for Black to move"),
        ("White root LOSS", "All twenty verified children are WIN for Black to move"),
        ("White root DRAW", "All twenty are resolved, none is LOSS for Black, and at least one is DRAW"),
        ("White root UNKNOWN", "Any required child is unverified or UNKNOWN, unless a verified LOSS child already proves WIN"),
    ]
    add_table(doc, ["Root value", "Complete quantified rule"], rows, widths=[1.55, 5.05], font_size=8.0, first_col_bold=True)
    doc.add_heading("10.1 Captured ledger state", level=2)
    status = validation["campaign"]
    rows = [
        ("Root obligations", status["root_obligations"]),
        ("Verified child values", status["verified_children"]),
        ("Root value", status["root_wdl"].upper()),
        ("Event hash chain", "valid" if status["event_chain_valid"] else "invalid"),
        ("Depth-4 leaf-path checkpoint", f"{status['depth4_leaf_paths']:,}"),
    ]
    add_table(doc, ["Metric", "Captured value"], rows, widths=[2.45, 4.15], font_size=8.4, first_col_bold=True)
    add_code(doc, "child values are from the child side-to-move perspective\none child LOSS -> root WIN\nall children WIN -> root LOSS\nresolved mix containing DRAW and no LOSS -> root DRAW\notherwise -> root UNKNOWN")
    add_callout(doc, "Current formal result", "The standard initial position is UNKNOWN in this package. No principal variation, engine score or opening preference is presented as a substitute.", fill="FFF4F2", accent=RED)
    add_page_break(doc)


def campaign(doc: Document) -> None:
    add_section_title(doc, "11", "Portable SQLite Proof Campaign", "Coordination state is mutable; proof records and the event journal are content-addressed and auditable.")
    add_pic(doc, "campaign_ledger.png", width=6.72, alt="SQLite campaign tables and event hash chain")
    caption(doc, "Jobs, events and metadata are separated; path references are package-relative so the campaign remains portable.")
    rows = [
        ("meta", "Root FEN/hashes, rule profile, schema and initialization data"),
        ("jobs", "20 obligations, lease owner/expiry, attempts, candidate WDL, certificate and checker paths/hashes"),
        ("events", "Append-only sequence, canonical payload, previous hash and event hash"),
        ("Lease", "BEGIN IMMEDIATE transaction selects one unresolved or expired job"),
        ("Candidate", "WIN/DRAW/LOSS plus certificate file hash; UNKNOWN cannot be submitted as a candidate"),
        ("Verification", "Separate JSON checker record must match obligation, value, certificate and child state"),
    ]
    add_table(doc, ["Element", "Contract"], rows, widths=[1.25, 5.35], font_size=7.8, first_col_bold=True)
    add_callout(doc, "Relocation test", "The database and its relative shard/certificate/checker paths were copied to another temporary location and reverified. This guards against absolute-path evidence that works only on the packaging machine.", fill="EAF8F2", accent=TEAL)
    add_page_break(doc)


def bounded_results(doc: Document) -> None:
    add_section_title(doc, "12", "Bounded Exact Results and Safe Incompleteness", "Two fixtures demonstrate existential closure and honest UNKNOWN propagation.")
    add_pic(doc, "bounded_results.png", width=6.72, alt="Exact mate-over-claim result and unknown initial depth two result")
    caption(doc, "The first fixture is exactly WIN despite an optional draw claim; the second remains UNKNOWN at the ply horizon.")
    rows = [
        ("Mate over claim", "8/8/8/1Q6/8/8/k7/2K5 w - - 100 2", "1", "2", "WIN", "Qa4# is an exact witness; draw claim is optional"),
        ("Initial position", "standard initial FEN", "2", "421", "UNKNOWN", "400 depth-2 leaf cutoffs remain open"),
    ]
    add_table(doc, ["Fixture", "Root", "Plies", "Nodes", "Value", "Reason"], rows, widths=[1.0, 2.0, 0.5, 0.65, 0.75, 1.7], font_size=6.9, first_col_bold=True)
    doc.add_heading("12.1 Certificate hashes", level=2)
    add_code(doc, "mate root certificate: 87d8a14d9b5af5cbfbccb9522f9c6ae372b66f68992c8749667a80dd1962dfb7\nmate child certificate: 891172e29d5aca1a49c857b4c2080284a98b717507e394061efd1f1f3ecb266f")
    add_callout(doc, "Proof lesson", "An existential WIN can close early when one exact losing child is found. Universal LOSS and complete DRAW cannot close until all relevant actions and history states are covered.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def mate_proof(doc: Document) -> None:
    add_section_title(doc, "13", "Retained Forced-Mate Certificate", "Chess 1.0's OR/AND mate proof remains a compact independently verifiable exact layer.")
    add_pic(doc, "mate_sequence.png", width=6.25, alt="Three-position sequence Qb5 Ka2 Qb2 checkmate")
    caption(doc, "Worked certificate: 1. Qb5 Ka2 2. Qb2# from the declared KQK fixture.")
    add_pic(doc, "proof_tree.png", width=6.15, alt="OR AND mate certificate tree")
    caption(doc, "The attacker's OR node supplies one winning move; the defender's AND node enumerates every legal reply.")
    add_callout(doc, "Exact scope", "This proves a forced mate from one serialized position within the declared horizon. It is not evidence about the game-theoretic value of the 32-piece initial position.", fill="EEF4FF", accent=BLUE)
    add_page_break(doc)


def tablebases(doc: Document) -> None:
    add_section_title(doc, "14", "Complete KQK and KRK WDL/DTM Tables", "Two finite three-piece state spaces are bundled as exact, reconstructible proof partitions.")
    add_pic(doc, "tablebase_counts.png", width=6.65, alt="KQK and KRK outcome counts")
    caption(doc, "All 2^19 packed addresses are classified as valid WIN/LOSS/DRAW or invalid placement.")
    rows = [
        ("KQK", f"{kq['valid_positions']:,}", f"{kq['outcome_counts']['win']:,}", f"{kq['outcome_counts']['loss']:,}", f"{kq['outcome_counts']['draw']:,}", kq["max_dtm_plies"], f"{kq['file_bytes']:,}"),
        ("KRK", f"{kr['valid_positions']:,}", f"{kr['outcome_counts']['win']:,}", f"{kr['outcome_counts']['loss']:,}", f"{kr['outcome_counts']['draw']:,}", kr["max_dtm_plies"], f"{kr['file_bytes']:,}"),
    ]
    add_table(doc, ["Table", "Valid", "WIN", "LOSS", "DRAW", "Max DTM", "gzip bytes"], rows, widths=[0.65, 1.0, 0.9, 0.9, 0.8, 0.75, 1.05], font_size=7.3, first_col_bold=True)
    add_callout(doc, "Tablebase boundary", "Outcome is from the side-to-move perspective. DTM is plies to checkmate under optimal play. These two tables are complete for their declared material signatures, not a substitute for all 3- through 7-piece endgames.", fill="EAF8F2", accent=TEAL)
    add_page_break(doc)


def tablebase_integration(doc: Document) -> None:
    add_section_title(doc, "15", "Finite Addressing, Generation and External Adapters", "Exact tables are proof partitions only when key, decode, legality and rule-profile contracts remain explicit.")
    rows = [
        ("Address width", "19 bits; 524,288 addresses for king, king, major piece and side-to-move fields"),
        ("Validity", "Overlapping pieces, adjacent kings and illegal check configurations are marked invalid"),
        ("Generation", "Retrograde fixed point seeds checkmates/stalemates and propagates side-to-move WDL/DTM"),
        ("Transport", "gzip payload plus JSON metadata, file length and SHA-256"),
        ("Probe", "Decode address, revalidate material/profile and return WDL/DTM"),
        ("External adapter", "May query larger Syzygy partitions, but must label rule profile and independently recheck transitions"),
    ]
    add_table(doc, ["Contract", "Exact requirement"], rows, widths=[1.45, 5.15], font_size=7.9, first_col_bold=True)
    doc.add_heading("15.1 Why these tables matter to a global attempt", level=2)
    add_bullets(doc, [
        "Solved material partitions become terminal proof nodes for larger predecessor graphs.",
        "Distance metadata supports deterministic witness extraction and cycle avoidance.",
        "Hashes permit distributed workers to refer to one exact table version rather than an ambiguous service response.",
        "A full campaign still needs promotion/capture transitions, 50/75-move policy compatibility and history-aware boundaries.",
    ], compact=False)
    add_callout(doc, "Compression rule", "Smaller storage is useful only with a total decoder and exact outcome equivalence. The tablebase key is a finite address, not a claim that chess state is one bit or that all history has disappeared.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def native_architecture(doc: Document) -> None:
    add_section_title(doc, "16", "Native C++20 Foundation", "A second implementation provides fast exact checks and a clean CUDA integration boundary.")
    add_pic(doc, "native_architecture.png", width=6.72, alt="C++20 modules and executable boundary")
    caption(doc, "Host legality and performance oracles are separate from the Python history/WDL authority.")
    rows = [
        ("core.cpp / fen.cpp", "Board representation, attacks, legal transitions, perft and notation transport"),
        ("search.cpp", "Deterministic alpha-beta and bounded mate search; scores are not global proof"),
        ("retrograde_cpu.cpp", "Generic monotone WIN/LOSS fixed-point demonstration"),
        ("sha256.cpp", "Dependency-free content hashing for native records"),
        ("ugts-chess2", "info, selftest, perft, search, mate, root-shards and retro-demo"),
        ("ugts-chess-gpu", "device-info, packed batch generation and CPU/CUDA self-test"),
    ]
    add_table(doc, ["Module", "Role"], rows, widths=[1.75, 4.85], font_size=7.7, first_col_bold=True)
    add_callout(doc, "Cross-implementation purpose", "The C++ layer is not merely a speed port. It provides a second executable implementation for perft, move proposals and fixed-point behavior, enabling differential checks before CUDA results are admitted.", fill="EEF4FF", accent=BLUE)
    add_page_break(doc)


def packed_protocol(doc: Document) -> None:
    add_section_title(doc, "17", "Packed Position and Move Proposal Protocol", "The wire format is simple enough to inspect, batch and compare across CPU and CUDA backends.")
    add_pic(doc, "gpu_protocol.png", width=6.72, alt="Packed position batch and independent Python comparison")
    caption(doc, "The simple kernel emits pseudo/legal candidates into a versioned binary batch; the Python oracle compares exact legal outputs.")
    rows = [
        ("PackedPosition", "64 bytes", "12 bitboards plus turn/rights/en-passant/clock/reserved fields"),
        ("Packed move", "16 bits", "6 source bits, 6 target bits, 3 promotion bits, spare version space"),
        ("Input file", "UGCB", "Header, record size/count and contiguous position records"),
        ("Output file", "UGMV", "Per-position counts plus up to 256 encoded moves each"),
        ("Ordering", "Deterministic", "Sorted/normalized before differential comparison"),
        ("Authority", "Proposal only", "Python exact legal moves remain the captured comparison oracle"),
    ]
    add_table(doc, ["Item", "Size/class", "Contract"], rows, widths=[1.35, 1.0, 4.25], font_size=7.7, first_col_bold=True)
    add_callout(doc, "Captured fallback", f"The host CPU backend processed {validation['packed_differential']['positions']} positions and {validation['packed_differential']['proposed_moves']} moves with zero mismatches. CUDA compilation and physical-device comparison remain pending.", fill="EAF8F2", accent=TEAL)
    add_page_break(doc)


def rtx_profile_page(doc: Document) -> None:
    add_section_title(doc, "18", "RTX 5070 Ti Laptop Execution Profile", "Official product identity is a starting constraint; runtime inspection is the authoritative device record.")
    rows = [
        ("Product", "GeForce RTX 5070 Ti Laptop GPU"),
        ("Architecture family", "NVIDIA Blackwell GeForce laptop"),
        ("CUDA cores", "5,888"),
        ("AI TOPS", "992"),
        ("Standard memory", "12 GB GDDR7"),
        ("CUDA compute capability", "12.0"),
        ("Native target", "sm_120 / compute_120"),
        ("Minimum package preset", "CUDA Toolkit 12.8 or newer; CUDA 12.8 added compiler support for SM120"),
    ]
    add_table(doc, ["Profile field", "Value"], rows, widths=[2.0, 4.6], font_size=8.2, first_col_bold=True)
    doc.add_heading("18.1 Runtime gate", level=2)
    add_code(doc, "nvidia-smi --query-gpu=name,memory.total,memory.free,compute_cap --format=csv\nugts-chess-gpu device-info\nnvcc --version")
    add_bullets(doc, [
        "Laptop power limits, clocks, free VRAM, cooling and driver/toolkit combinations vary by model and power mode.",
        "The build preset targets SM120. Datacenter Blackwell SM100a architecture-conditional kernels are not assumed compatible with GeForce RTX 50-series SM120.",
        "The package records the actual device name, compute capability, memory, SM count, warp size and maximum threads before a benchmark is promoted.",
    ], compact=False)
    add_source_note(doc, "Official NVIDIA GeForce RTX 50 Series laptop specifications, CUDA GPU Compute Capability mapping, CUDA 12.8 release notes, Blackwell compatibility guide and CUTLASS architecture note.")
    add_page_break(doc)


def memory_plan(doc: Document) -> None:
    add_section_title(doc, "19", "VRAM Budget and Out-of-Core Design", "The nominal 12 GB frame buffer is a bounded cache, not a complete home for the chess proof graph.")
    add_pic(doc, "rtx_memory.png", width=6.72, alt="Nine GiB solver allocation and three GiB safety headroom")
    caption(doc, "The checked-in starting profile reserves 3 GiB and assigns 9 GiB to solver structures.")
    alloc = rtx_profile["allocation_mib"]
    rows = [
        ("Proof/transposition index", f"{alloc['transposition_and_proof_index']:,} MiB", "Content-addressed lookup and verified-state cache"),
        ("Frontier", f"{alloc['frontier_positions']:,} MiB", "Current batch of unresolved positions"),
        ("Move buffers", f"{alloc['move_matrix_and_counts']:,} MiB", "Counts and 16-bit proposal arrays"),
        ("Retrograde", f"{alloc['retrograde_edges_and_counters']:,} MiB", "Outcome vectors and graph work arrays"),
        ("Checkpoint", f"{alloc['checkpoint_staging']:,} MiB", "Device-to-host export and integrity staging"),
        ("Scratch", f"{alloc['scratch']:,} MiB", "Temporary scans, prefix sums and error safety"),
        ("Reserved headroom", f"{rtx_profile['reserved_headroom_mib']:,} MiB", "Driver/display/fragmentation/thermal fallback margin"),
    ]
    add_table(doc, ["Allocation", "Starting size", "Purpose"], rows, widths=[1.75, 1.2, 3.65], font_size=7.7, first_col_bold=True)
    add_callout(doc, "Capacity reality", "Even 9 GiB divided by the 64-byte exchange record holds at most about 151 million records before indexes, move arrays, histories and certificates. A proof-scale campaign must stream, partition, checkpoint and persist to disk.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def cuda_boundary(doc: Document) -> None:
    add_section_title(doc, "20", "CUDA Kernels and Correctness Boundary", "The first GPU code is deliberately simple so that equal-output validation comes before optimization.")
    rows = [
        ("Scalar batch expander", "One packed position per thread; fixed 256-move slot; output compared with host oracle"),
        ("Generic retrograde step", "One graph node per thread; monotone UNKNOWN -> WIN/LOSS propagation"),
        ("Device inspector", "Queries name, compute capability, total/free memory, SMs, warp and block limits"),
        ("CPU fallback", "Runs the identical batch protocol when CUDA is absent or rejected"),
        ("CMake profile", "UGTS_ENABLE_CUDA=ON; CMAKE_CUDA_ARCHITECTURES=120"),
        ("Promotion gate", "No mismatches, deterministic record order, full reconstructibility and independent proof checker"),
    ]
    add_table(doc, ["Component", "Current implementation"], rows, widths=[1.75, 4.85], font_size=7.8, first_col_bold=True)
    doc.add_heading("20.1 Optimization sequence after correctness", level=2)
    add_numbered(doc, [
        "Measure scalar kernel occupancy, transfer overhead and batch-size behavior on the actual laptop.",
        "Add warp-cooperative attack generation and prefix-sum compaction only behind the same binary protocol.",
        "Separate hot structure-of-arrays data from the 64-byte exchange record while retaining a total versioned decoder.",
        "Add asynchronous streams and double buffering; preserve deterministic per-position move ordering at the comparison boundary.",
        "Keep WDL certificate checking on an independent host path until a separately verified GPU checker exists.",
    ])
    add_callout(doc, "Packaging evidence", "The delivery host had no nvcc and no physical RTX 5070 Ti. CUDA throughput, VRAM, power, temperature, clocks and long-run stability are intentionally unclaimed.", fill="FFF4F2", accent=RED)
    add_page_break(doc)


def codex_workflow(doc: Document) -> None:
    add_section_title(doc, "21", "Codex Build, Measure and Promotion Workflow", "The supplied scripts preserve device evidence under validation/device rather than overwriting host evidence.")
    add_pic(doc, "codex_workflow.png", width=6.72, alt="Six-stage Codex workflow for building and promoting RTX results")
    caption(doc, "Device capture, SM120 build, differential testing, measurement, proof-DAG extension and shard closure are separate gates.")
    doc.add_heading("21.1 First commands on Windows", level=2)
    add_code(doc, "powershell -ExecutionPolicy Bypass -File scripts/build_rtx5070ti.ps1\npowershell -ExecutionPolicy Bypass -File scripts/run_codex_campaign.ps1")
    doc.add_heading("21.2 Required measurements", level=2)
    add_bullets(doc, [
        "p50/p95/p99 batch latency, positions/s, moves/s and transfer/compute breakdown",
        "peak VRAM, host RAM and storage bytes per verified node",
        "5-, 15- and 30-minute clocks, temperature, power, fallback and throttling behavior",
        "differential mismatch count across packaged fixtures and a seeded random legal corpus",
        "checkpoint/restart reproducibility and event-chain validation after process interruption",
    ], compact=False)
    add_callout(doc, "Release blocker", "Any extra move, missing move, history ambiguity, incomplete AND-node reply set, corrupt checkpoint, hash mismatch or device-dependent result blocks proof promotion.", fill="FFF4F2", accent=RED)
    add_page_break(doc)


def global_architecture(doc: Document) -> None:
    add_section_title(doc, "22", "Path from Laptop Experiment to Global Proof Campaign", "Progress is measured by independently verifiable obligation closure, not by engine rating or searched nodes alone.")
    add_pic(doc, "roadmap.png", width=6.72, alt="Roadmap from frozen authority through initial promotion")
    caption(doc, "The initial node is promoted only after proof infrastructure, exact coverage and checker diversity are complete.")
    doc.add_heading("22.1 Disk-backed proof DAG", level=2)
    add_bullets(doc, [
        "Partition by material signature, rights/history class and canonical hash prefix; store full reconstructible state beside any short index.",
        "Use append-only frontier and certificate objects with schema IDs, CRC/SHA checks and content-addressed child references.",
        "Treat opening books, engines and neural policies as prioritizers. They cannot close a proof node.",
        "Integrate exact tablebase partitions through profile-labeled adapters and boundary transition checks.",
        "Merge root shards only after independent checker records agree; retain conflicts and rejected candidates in lineage.",
    ], compact=False)
    add_callout(doc, "Scale boundary", "A single 12 GB laptop can validate kernels, close small partitions and contribute shards. It is not represented as sufficient by itself to exhaust the full classical game graph.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def validation_page(doc: Document) -> None:
    add_section_title(doc, "23", "Captured Host Validation", "The supplied evidence establishes exact host behavior and packaging integrity, not physical-GPU performance.")
    add_pic(doc, "validation_dashboard.png", width=6.72, alt="Host validation dashboard")
    caption(doc, "All captured host gates passed; the CUDA device lane remains explicitly pending.")
    rows = [
        ("Python tests", validation["python_tests"]["count"], "PASS"),
        ("Python perft checks", validation["python_perft"]["checks"], "PASS"),
        ("C++ CTest", validation["native_build"]["ctest_total"], "PASS"),
        ("Native packed self-tests", validation["native_selftests"]["packed"]["passed"], "PASS"),
        ("Schema documents", validation["schemas_and_proofs"]["schema_documents"], "PASS"),
        ("Mechanism records", validation["mechanism_catalog"]["entries"], "PASS"),
        ("Wheel clean install", validation["wheel"]["file"], "PASS"),
        ("Physical RTX run", 0, "PENDING"),
    ]
    add_table(doc, ["Gate", "Count/artifact", "Status"], rows, widths=[2.35, 2.65, 1.6], font_size=8.0, first_col_bold=True)
    add_callout(doc, "Validation meaning", validation["claim_boundary"], fill="EEF4FF", accent=BLUE)
    add_page_break(doc)


def performance_page(doc: Document) -> None:
    add_section_title(doc, "24", "Performance Evidence and Non-Claims", "Host timing is useful for regression but cannot be generalized to the user's laptop before measurement.")
    rows = [
        ("Native initial perft depth 5", f"{native_depth5['nodes']:,} nodes", f"{native_depth5['seconds']:.3f} s", f"{native_depth5['nodes_per_second']/1_000_000:.3f} Mnodes/s"),
        ("Maximum RSS during captured depth-5 run", "1,800 KiB", "Linux container", "host-specific"),
        ("Packed differential", f"{validation['packed_differential']['positions']} positions / {validation['packed_differential']['proposed_moves']} moves", "0 mismatches", "CPU fallback"),
        ("Tablebase transport", f"{validation['tablebases']['compressed_bytes']:,} bytes", "2 exact files", "hash verified"),
        ("Python wheel", f"{validation['wheel']['bytes']:,} bytes", "isolated install", "PASS"),
    ]
    add_table(doc, ["Measurement", "Work", "Observed", "Boundary"], rows, widths=[2.05, 1.8, 1.25, 1.5], font_size=7.55, first_col_bold=True)
    doc.add_heading("24.1 What is not claimed", level=2)
    add_bullets(doc, [
        "No Stockfish, Leela Chess Zero or other mature-engine strength comparison.",
        "No RTX speedup, positions/s, energy efficiency, thermals, battery life or sustained laptop throughput.",
        "No implication that faster move generation changes the mathematical WDL obligations.",
        "No universal bit-identical floating-point behavior; promoted legality, hashes and proof records use exact integer/serialized boundaries.",
        "No proof that the Python interpreter, C++ compiler, operating system or GPU hardware is free of defects.",
    ], compact=False, accent=RED)
    add_callout(doc, "Equal-output rule", "A performance claim becomes relevant only after CPU and CUDA outputs agree under the same rule profile and checker. Correctness is not inferred from speed.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def roadmap_kill(doc: Document) -> None:
    add_section_title(doc, "25", "Promotion Gates, Roadmap and Kill Criteria", "The campaign is designed to fail safely and visibly when evidence is incomplete.")
    doc.add_heading("25.1 Promotion sequence", level=2)
    add_numbered(doc, [
        "Freeze legal and history authority; maintain cross-implementation perft and transition tests.",
        "Compile and differential-test SM120 kernels on the target laptop.",
        "Add append-only frontier/proof storage and crash-safe checkpoints.",
        "Expand exact endgame partitions with rule-profile and hash commitments.",
        "Implement complete dead-position certificates or preserve UNKNOWN.",
        "Close one root shard at a time; diversify independent checkers.",
        "Promote the initial node only when its complete quantified WDL rule is satisfied.",
    ])
    doc.add_heading("25.2 Release-blocking failures", level=2)
    add_bullets(doc, [
        "Any legal-move or terminal mismatch, including castling, en-passant, promotion, check or draw-order errors.",
        "A proof verifier accepts a missing defender reply, illegal edge, wrong pre-state, mismatched hash or unsupported draw closure.",
        "A 64-bit key, score, GPU hash or cache hit is treated as complete proof identity.",
        "Repetition history, draw-claim ownership, dead-position scope or 50/75-move policy is implicit.",
        "A worker directly changes authoritative WDL without an independent checker record.",
        "Compression or seed reconstruction discards information needed to reproduce legal actions or certificates.",
        "The UGTS wrapper becomes less auditable than the conventional exact rule/proof implementation it is meant to discipline.",
    ], compact=False, accent=RED)
    add_callout(doc, "Stop condition", "When a boundary cannot be certified, the correct output is UNKNOWN plus a reason-coded open obligation. That is a successful safety behavior, not a reason to fabricate a value.", fill="F4F0FF", accent=PURPLE)
    add_page_break(doc)


def catalog(doc: Document) -> None:
    add_section_title(doc, "26", "Chess Mechanism Catalog C001-C104", "Namespaced mechanism records provide traceability without renumbering the broader UGTS catalogs.")
    counts = Counter(item["domain"] for item in mechanisms)
    rows = []
    for domain, count in sorted(counts.items(), key=lambda item: (-item[1], item[0])):
        ids = [item["id"] for item in mechanisms if item["domain"] == domain]
        names = [item["mechanism"] for item in mechanisms if item["domain"] == domain]
        rows.append((domain, count, f"{ids[0]}-{ids[-1]}", "; ".join(names[:3]) + ("; ..." if len(names) > 3 else "")))
    add_table(doc, ["Domain", "Count", "ID span", "Representative records"], rows, widths=[1.25, 0.55, 0.9, 3.9], font_size=6.9, first_col_bold=True)
    doc.add_heading("26.1 Representative 2.0 additions", level=2)
    selected_ids = {"C065", "C069", "C073", "C078", "C082", "C087", "C092", "C097", "C101", "C104"}
    selected = [item for item in mechanisms if item["id"] in selected_ids]
    rows = [(item["id"], item["domain"], item["mechanism"], item["definition"]) for item in selected]
    add_table(doc, ["ID", "Domain", "Mechanism", "Normalized definition"], rows, widths=[0.45, 1.0, 1.45, 3.7], font_size=6.65, first_col_bold=True)
    add_callout(doc, "Catalog meaning", "An ID records scope, implementation and validation linkage. It is not by itself a claim of scientific novelty, authorship priority, patentability or proof completeness.", fill="FFF7E8", accent=GOLD)
    add_page_break(doc)


def package_map(doc: Document) -> None:
    add_section_title(doc, "27", "Accompanying ZIP and Reproduction", "The archive is a complete Codex handoff: source, exact data, schemas, campaign state, build scripts, report and captured evidence.")
    add_pic(doc, "package_map.png", width=6.72, alt="Package folders for Python C++ CUDA campaign data specifications scripts and validation")
    caption(doc, "Source PDFs are referenced by title, page basis and SHA-256 but are not redistributed.")
    rows = [
        ("src/ugts_chess", "Python rule/history oracle, WDL certificates, campaign, tablebases and CLI"),
        ("cpp", "C++20 core, host tests, packed protocol, CPU fallback and optional CUDA kernels"),
        ("examples/campaign", "SQLite root campaign, 20 shard JSON records and exact depth-4 workload checkpoint"),
        ("data", "KQK/KRK compressed exact tables and metadata"),
        ("spec", "Formal definition, schemas, source register, RTX profile and C001-C104 catalog"),
        ("scripts", "Codex task, Windows SM120 build and campaign run scripts"),
        ("validation", "Host tests, CTest, differential outputs, wheel, manifests and device placeholder"),
        ("report / web", "Editable DOCX, final PDF, diagrams and offline mate-proof viewer"),
    ]
    add_table(doc, ["Path", "Contents"], rows, widths=[1.65, 4.95], font_size=7.75, first_col_bold=True)
    doc.add_heading("27.1 Host reproduction", level=2)
    add_code(doc, "cd UGTS_KC_CHESS_2_0_RTX5070TI\nPYTHONPATH=src python -m unittest discover -s tests -v\nbash scripts/build_host.sh\nPYTHONPATH=src python tools/run_validation_v2.py\npython tools/generate_report_assets.py\npython tools/build_report.py")
    add_page_break(doc)


def final_definition(doc: Document) -> None:
    add_section_title(doc, "28", "Final Upgraded Definition", "The foundation is a proof campaign with optional accelerators, not a search score presented as a solved game.")
    add_callout(
        doc,
        "Formal definition",
        "UGTS Chess Proof Campaign 2.0 is a finite, content-addressed application profile in which complete classical positions and exact repetition context define legal actions; special rights and king safety certify events; immutable transitions update state; claimable draws remain player actions; automatic terminals seed four-valued WDL obligations; certificates record existential witnesses or complete universal coverage; an independent checker promotes results into a hash-chained campaign ledger; exact KQK/KRK tables close bounded partitions; and C++20/CUDA adapters accelerate proposals and fixed-point work without becoming final proof authority.",
        fill="EAF8F2",
        accent=TEAL,
    )
    doc.add_heading("28.1 Canonical object", level=2)
    add_code(doc, "UGTS-CHESS-2 = (Q, A, M, G, T, H, V, C, P, L, X)\nQ complete typed state         A attack/support relations\nM deterministic legal moves   G terminal and draw guards\nT immutable transitions       H exact history context\nV WIN|LOSS|DRAW|UNKNOWN       C certificates/checkers\nP campaign partitions         L replay/proof lineage\nX optional CPU/CUDA adapters")
    doc.add_heading("28.2 Current root theorem", level=2)
    add_code(doc, "initial_position_value = UNKNOWN\nreason = 20 exact root obligations exist; 0 child WDL certificates have been independently promoted")
    add_callout(doc, "Release interpretation", "The package materially advances an attempt by making the open proof explicit, partitioned, portable and executable. It does not overstate that progress as a game-theoretic solution.", fill="EEF4FF", accent=BLUE)
    add_page_break(doc)


def references(doc: Document) -> None:
    add_section_title(doc, "Appendix", "Source Register, External References and Attribution", "Supplied artifacts and current standards are separated from engineering additions and captured package evidence.")
    rows = []
    for index, source in enumerate(source_reg["sources"], start=1):
        rows.append((f"S{index}", source["title"], source["page_basis"], source["sha256"][:16] + "...", source["role"]))
    add_table(doc, ["ID", "Supplied source", "Page basis", "SHA-256", "Use"], rows, widths=[0.35, 1.35, 0.85, 1.15, 2.9], font_size=5.95, first_col_bold=True)
    doc.add_heading("External standards and current device references", level=2)
    ext_rows = []
    for index, item in enumerate(source_reg["external_standards"], start=1):
        ext_rows.append((f"E{index}", item["title"], item.get("edition", "current online source"), item["role"], item["url"]))
    add_table(doc, ["ID", "Reference", "Edition", "Use", "Location"], ext_rows, widths=[0.35, 1.35, 1.0, 2.25, 1.65], font_size=5.75, first_col_bold=True)
    doc.add_heading("Attribution and legal boundary", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Prepared for Tom Klootwijk. ")
    r.bold = True
    p.add_run("The requester attribution is recorded as supplied and has not been independently verified. This report and package are technical design and validation artifacts. They are not legal proof of identity, authorship, ownership, priority, patentability, exclusive rights, licensing status or scientific consensus.")
    add_callout(doc, "Closing statement", "The strongest evidence in this release is exact and local: legal transitions, bounded certificates, finite tablebases, campaign hashes and passing reproducibility gates. The global classical root remains UNKNOWN until the missing proof obligations are actually closed.", fill="F4F0FF", accent=PURPLE)


def main() -> None:
    REPORT.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    cover(doc)
    executive(doc)
    contents(doc)
    source_basis(doc)
    solution_meaning(doc)
    fide_profile(doc)
    authority(doc)
    state_identity(doc)
    draw_history(doc)
    wdl(doc)
    certificates(doc)
    root_obligations(doc)
    root_aggregation(doc)
    campaign(doc)
    bounded_results(doc)
    mate_proof(doc)
    tablebases(doc)
    tablebase_integration(doc)
    native_architecture(doc)
    packed_protocol(doc)
    rtx_profile_page(doc)
    memory_plan(doc)
    cuda_boundary(doc)
    codex_workflow(doc)
    global_architecture(doc)
    validation_page(doc)
    performance_page(doc)
    roadmap_kill(doc)
    catalog(doc)
    package_map(doc)
    final_definition(doc)
    references(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
